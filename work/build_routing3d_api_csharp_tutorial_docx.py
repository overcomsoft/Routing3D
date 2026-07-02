from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Routing3D_API_Manual_CSharp_Tutorial_20260702.docx"

FONT_KR = "Malgun Gothic"
FONT_CODE = "Consolas"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D9DEE7"
TEXT = "1F2937"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size: float = 10.0, bold: bool | None = None, color: str | None = None,
                 font: str = FONT_KR) -> None:
    run.font.name = font
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:eastAsia"), FONT_KR)


def add_para(doc: Document, text: str = "", size: float = 10.0, bold: bool = False,
             color: str | None = TEXT, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(11 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 12.5, bold=True, color=BLUE)


def add_note(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF4EC")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "  ")
    set_run_font(r, size=9.5, bold=True, color="20603D")
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=TEXT)


def add_code(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for idx, line in enumerate(code.strip("\n").splitlines()):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        set_run_font(r, size=8.2, color="111827", font=FONT_CODE)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_BLUE)
        if widths:
            set_cell_width(hdr[i], widths[i])
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=8.8, bold=True, color="123A5A")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if widths:
                set_cell_width(cells[i], widths[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            for j, line in enumerate(value.split("\n")):
                if j:
                    p.add_run("\n")
                r = p.add_run(line)
                set_run_font(r, size=8.5, color=TEXT, font=FONT_CODE if "`" in line else FONT_KR)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    styles = doc.styles
    styles["Normal"].font.name = FONT_KR
    styles["Normal"].font.size = Pt(10)
    return doc


def add_cover(doc: Document) -> None:
    add_para(doc, "Routing3D API Manual", size=22, bold=True, color=BLUE)
    add_para(doc, "C++ C ABI 및 C# 사용 튜토리얼", size=17, bold=True, color="374151")
    add_para(doc, "작성일: 2026-07-02", size=10, color="4B5563")
    add_para(doc, "대상: routing3d_capi.dll, cpp/capi/routing3d_capi.h, csharp/Routing3D.Engine", size=9.5, color="4B5563")
    add_note(doc, "범위", "현재 저장소에 구현된 네이티브 API와 C# Managed wrapper 기준입니다. 모든 좌표와 길이는 mm 단위이며, 경로 셀은 I/J/K 정수 좌표입니다.")
    doc.add_paragraph()
    add_table(doc, ["문서 구성", "내용"], [
        ["1. 빠른 시작", "C# 프로젝트에서 DLL 배치, Routing3DEngine 생성, 단일/다중 라우팅 실행"],
        ["2. API 매뉴얼", "C ABI 구조체, 생명주기, 입력, 옵션, 라우팅 함수, 결과 복사, 진단 API"],
        ["3. C# 튜토리얼", "scene.txt 라우팅, 직접 핸들 방식, 진행률/취소, 고급 옵션"],
        ["4. 운영 체크리스트", "실패 사유, 성능 옵션, 검증 절차"],
    ], [1.7, 5.8])
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_architecture(doc: Document) -> None:
    add_heading(doc, "1. 전체 구조", 1)
    add_para(doc, "Routing3D는 C++ 코어 엔진을 C ABI로 감싸고, C#에서는 P/Invoke와 SafeHandle 기반 래퍼로 호출합니다.")
    add_code(doc, """
C# WPF / .NET app
  -> Routing3D.Engine.Routing3DEngine
    -> Native.cs P/Invoke (cdecl, x64, UTF-8)
      -> routing3d_capi.dll
        -> routing3d core: occupancy, A*, multi_route, corridor, octree diagnostics
""")
    add_table(doc, ["레이어", "주요 파일", "역할"], [
        ["C++ core", "cpp/include/routing3d/*.hpp", "격자, 점유맵, A*, multi-route, corridor, scene I/O"],
        ["C ABI", "cpp/capi/routing3d_capi.h/.cpp", "외부 호출용 안정 ABI. 예외를 R3dStatus로 변환"],
        ["C# interop", "csharp/Routing3D.Engine/Native.cs", "DllImport, blittable struct, UTF-8 문자열 처리"],
        ["C# wrapper", "csharp/Routing3D.Engine/Routing3DEngine.cs", "IDisposable 엔진 핸들, 객체 모델, 결과 변환"],
        ["Viewer/diagnostics", "csharp/Routing3D.Viewer", "DB 실데이터 로딩, 진행률, 진단, 가시화"],
    ], [1.3, 2.5, 3.7])


def add_quick_start(doc: Document) -> None:
    add_heading(doc, "2. C# 빠른 시작", 1)
    add_para(doc, "가장 단순한 사용 패턴은 엔진 생성, 격자/파라미터 설정, 장애물/작업 추가, RouteMulti 실행, GetResult 조회입니다.")
    add_code(doc, """
using Routing3D.Engine;

using var engine = new Routing3DEngine();

engine.SetGrid(cellMm: 50, ox: 0, oy: 0, oz: 0, nx: 200, ny: 160, nz: 40);
engine.SetParameters(new RoutingParameters
{
    CellMm = 50,
    TurnCostMm = 500,
    ClearanceCostMm = 10,
    ClearanceRadiusCells = 2,
    ClearanceConnectivity = 6
});

engine.AddObstacle(2000, 2000, 0, 3000, 3000, 2000);
int task = engine.AddTask(100, 100, 500, 9000, 7000, 500, "ACID", "Exhaust");

engine.RouteMulti("longest");
RouteResult result = engine.GetResult(task);

Console.WriteLine($"{result.Success}, len={result.LengthMm}, turns={result.Turns}");
foreach (PathCell c in result.Path)
    Console.WriteLine($"{c.I}, {c.J}, {c.K}");
""")
    add_note(doc, "배포", "C# 실행 파일과 같은 폴더 또는 PATH에 routing3d_capi.dll이 있어야 합니다. x64 프로세스에서 호출해야 합니다.")


def add_c_api(doc: Document) -> None:
    add_heading(doc, "3. C ABI API 매뉴얼", 1)
    add_heading(doc, "3.1 기본 구조체와 상태 코드", 2)
    add_table(doc, ["항목", "필드/값", "설명"], [
        ["R3dStatus", "OK=0, ERR_ARG=1, ERR_PARSE=2, ERR_RUNTIME=3, ERR_RANGE=4", "모든 C ABI 함수의 공통 상태 코드"],
        ["R3dGrid", "cell_mm, ox, oy, oz, nx, ny, nz", "격자 해상도, 원점, 크기"],
        ["R3dParams", "cell_mm, w_turn, w_clear, w_corridor, w_heur, w_heur_near, clearance_radius, clearance_connectivity, corridor_radius, rack_levels", "A* 비용과 탐색 옵션"],
        ["R3dResult", "success, length_mm, cost_mm, turns, expanded_nodes, elapsed_ms, path_len, visited_len, fail_reason", "작업 1건의 결과와 복사할 배열 길이"],
        ["R3dOctreeLeaf", "x0_mm, y0_mm, z0_mm, size_mm, state", "Octree 진단/가시화용 leaf"],
    ], [1.4, 3.2, 3.0])
    add_heading(doc, "3.2 생명주기와 입력", 2)
    add_table(doc, ["함수", "용도", "C# wrapper"], [
        ["r3d_create / r3d_destroy", "엔진 핸들 생성/해제", "Routing3DEngine 생성자 / Dispose"],
        ["r3d_load_scene_text", "scene.txt 전체 로드", "LoadSceneText"],
        ["r3d_set_grid", "격자 정의", "SetGrid"],
        ["r3d_set_params", "비용/탐색 파라미터", "SetParameters"],
        ["r3d_add_obstacle", "충돌 AABB 추가", "AddObstacle"],
        ["r3d_add_passthrough", "통과 객체 AABB 추가", "AddPassthrough"],
        ["r3d_add_task", "라우팅 작업 추가 후 task index 반환", "AddTask"],
        ["r3d_set_task_endpoints", "기존 작업의 시작/목표 갱신", "SetTaskEndpoints"],
        ["r3d_set_task_diameter", "관경 기반 반경/우선순위 입력", "SetTaskDiameter"],
        ["r3d_set_task_goal_dir", "목표 진입축 제약", "SetTaskGoalDirection"],
    ], [2.2, 3.4, 1.9])
    add_heading(doc, "3.3 라우팅 함수", 2)
    add_table(doc, ["함수", "설명", "권장 사용"], [
        ["r3d_route_scene_text", "scene.txt 문자열을 한 번에 라우팅하고 결과 scene.txt 반환", "간단한 CLI/테스트"],
        ["r3d_route_multi", "전체 작업을 priority 순서로 순차 라우팅", "일반 배치 라우팅 기본"],
        ["r3d_route_multi_progress", "RouteMulti + 진행률/취소 콜백", "GUI, 장시간 DB 실데이터"],
        ["r3d_route_task", "단일 작업을 원본 장애물 기준으로 탐색", "부분 재탐색, 진단"],
        ["r3d_route_ripup", "실패 작업 기준으로 blocker를 걷어내며 재시도", "혼잡/교착 회복"],
        ["r3d_route_corridor / r3d_route_corridor_multi", "coarse-to-fine corridor 방식", "초대형 격자 실험/특정 회랑 유도"],
        ["r3d_route_task_octree", "Octree leaf 기반 단일 작업 탐색", "현재는 진단/옵트인 성격"],
    ], [2.4, 3.6, 1.7])
    add_heading(doc, "3.4 옵션 setter", 2)
    add_table(doc, ["함수", "기본값", "효과"], [
        ["r3d_set_collect_visited", "on", "방문 셀 수집. 성능/메모리가 중요하면 off"],
        ["r3d_set_pipe_radius", "0", "모든 배관에 동일 셀 반경 적용"],
        ["r3d_set_per_task_radius", "off", "task별 diameter에서 셀 반경 계산"],
        ["r3d_set_pipe_gap", "0 mm", "배관 간 최소 이격을 반경 팽창에 반영"],
        ["r3d_set_cbs_depth", "0", "CBS-lite 충돌 회복 깊이"],
        ["r3d_set_min_straight / r3d_set_min_straight_mm", "0", "코너 주변 최소 직선 길이 제약"],
        ["r3d_set_corridor_cells", "none", "학습/기존설계 회랑 셀 주입"],
        ["r3d_set_segment_astar / r3d_set_octree_guide / r3d_set_route_split", "off", "대형 경로 실험용 옵션"],
    ], [2.5, 1.1, 3.7])
    add_heading(doc, "3.5 결과 복사와 문자열 소유권", 2)
    add_table(doc, ["API", "규칙"], [
        ["r3d_get_result", "task별 R3dResult 조회. path_len/visited_len으로 복사 버퍼 크기를 결정"],
        ["r3d_copy_path", "int32 배열에 I,J,K 순서로 복사. 배열 크기는 3 * path_len"],
        ["r3d_copy_visited", "A* 확장 셀 복사. collect_visited가 꺼져 있으면 0일 수 있음"],
        ["r3d_copy_blocked / r3d_copy_passthrough", "점유맵/통과 객체 셀을 복사하여 가시화에 사용"],
        ["r3d_dump_scene_text / r3d_get_runtime_report", "char** 문자열 반환. 호출자는 r3d_free_string으로 해제"],
    ], [2.4, 5.0])


def add_csharp_wrapper(doc: Document) -> None:
    add_heading(doc, "4. C# Managed Wrapper 사용법", 1)
    add_table(doc, ["C# 타입/메서드", "설명"], [
        ["Routing3DEngine", "IDisposable 엔진 핸들. using 블록으로 수명 관리"],
        ["RoutingGrid, RoutingParameters", "격자와 비용 파라미터 모델"],
        ["Aabb, Vec3, PathCell", "mm 좌표/AABB와 IJK 경로 셀"],
        ["RouteResult", "Success, LengthMm, CostMm, Turns, ExpandedNodes, ElapsedMs, Fail, Path, Visited"],
        ["RouteProgress", "RouteMultiProgress 콜백 이벤트 모델"],
        ["RouteFail", "None, StartBlocked, GoalBlocked, CorridorMiss, ExpansionLimit, GoalDirBlocked, NoPath"],
        ["GoalDirection", "Any 또는 6-neighbor 진입축"],
    ], [2.4, 5.1])
    add_heading(doc, "4.1 scene.txt 한 번에 라우팅", 2)
    add_code(doc, """
string sceneText = File.ReadAllText("input.scene.txt", Encoding.UTF8);
string routed = Routing3DEngine.RouteSceneText(sceneText, mode: "multi", priority: "longest");
File.WriteAllText("output.scene.txt", routed, Encoding.UTF8);
""")
    add_heading(doc, "4.2 다중 배관 + 관경/이격 옵션", 2)
    add_code(doc, """
using var engine = new Routing3DEngine();
engine.SetGrid(new RoutingGrid(50, new Vec3(-1000, -1000, 0), 500, 500, 80));
engine.SetParameters(new RoutingParameters
{
    CellMm = 50,
    TurnCostMm = 500,
    ClearanceCostMm = 10,
    HeuristicWeight = 1.0,
    ClearanceRadiusCells = 2
});

foreach (Aabb obstacle in obstacles)
    engine.AddObstacle(obstacle);

int acid = engine.AddTask(startAcid, endDuct, "ACID", "Exhaust");
int alk  = engine.AddTask(startAlk,  endDuct, "ALKA", "Exhaust");
engine.SetTaskDiameter(acid, 150);
engine.SetTaskDiameter(alk, 100);

engine.SetPerTaskRadius(true);
engine.SetPipeGap(60);
engine.SetCbsDepth(2);
engine.SetMinStraight(2.0);

engine.RouteMulti("diameter");
RouteResult acidResult = engine.GetResult(acid);
""")
    add_heading(doc, "4.3 진행률 표시와 취소", 2)
    add_code(doc, """
var cancel = false;
engine.RouteMultiProgress("longest", p =>
{
    if (p.Phase == 0)
        Console.WriteLine($"search task={p.TaskIndex}, {p.Progress01:P0}, expanded={p.ExpandedNodes}");
    else
        Console.WriteLine($"done {p.Done}/{p.Total}, success={p.Success}, len={p.LengthMm}");
}, shouldCancel: () => cancel);
""")
    add_heading(doc, "4.4 런타임 리포트와 진단", 2)
    add_code(doc, """
engine.SetCollectVisited(true);
engine.RouteMulti("longest");

string reportJson = engine.GetRuntimeReportJson();
IReadOnlyList<PathCell> blocked = engine.CopyBlocked();
IReadOnlyList<PathCell> passthrough = engine.CopyPassthrough();
RouteResult r = engine.GetResult(taskIndex);

if (!r.Success)
    Console.WriteLine($"fail={r.Fail}, expanded={r.ExpandedNodes}, visited={r.Visited.Count}");
""")


def add_operational_guidance(doc: Document) -> None:
    add_heading(doc, "5. 실데이터 운영 기준", 1)
    add_table(doc, ["상황", "우선 확인", "권장 대응"], [
        ["StartBlocked / GoalBlocked", "PoC가 장비/장애물 AABB 내부 또는 경계에 있는지", "PoC surface snap, start/goal 보정, blocked map 표시"],
        ["ExpansionLimit", "cell_mm, R3D_MAX_EXP, expanded_nodes", "cell 50/100 비교, weighted heuristic, corridor 후보 검토"],
        ["NoPath", "blocked/passthrough map, pipe radius/gap, CBS depth", "gap/반경 완화 A/B, rip-up/CBS, 라우팅 순서 변경"],
        ["GUI가 오래 멈춤", "RouteMultiProgress 사용 여부, collect_visited, cell 크기", "백그라운드 worker, visited off, cell 50 이상 실측"],
        ["DLL mismatch", "Viewer bin의 routing3d_capi.dll 시간/크기", "C++ Release 빌드 후 실행 폴더에 복사"],
    ], [1.8, 2.7, 3.0])
    add_heading(doc, "5.1 성능 옵션 판단", 2)
    add_table(doc, ["옵션", "효과", "주의"], [
        ["cell_mm 증가", "노드 수를 3차원으로 줄여 가장 강력한 속도 개선", "좁은 통로 해상도 손실 가능"],
        ["SetCollectVisited(false)", "방문 셀 저장 메모리 감소", "탐색 애니메이션/방문맵 사용 불가"],
        ["HeuristicWeight > 1", "확장 노드 감소 가능", "최단/골든 불변성보다 속도 우선"],
        ["SetCbsDepth", "혼잡 경로 회복률 향상", "분기 증가로 케이스에 따라 느려질 수 있음"],
        ["Octree/Segment/RouteSplit", "대형 격자 실험 옵션", "현재 Exhaust 실데이터에서는 기본 A*보다 빠르다고 증명되지 않음"],
    ], [2.0, 2.9, 2.6])
    add_heading(doc, "5.2 검증 체크리스트", 2)
    add_table(doc, ["체크", "명령/방법"], [
        ["C++ ABI 회귀", "ctest --test-dir cpp/build -C Release -R \"capi|golden|octree|implicit\""],
        ["C# 빌드", "dotnet build csharp/Routing3D.Viewer.sln -c Release"],
        ["DB 실데이터 재현", "Routing3D.Viewer.exe --dbroute <projectId> <cellMm> <scope> <outPath>"],
        ["결과 품질", "성공률, total length, turns, expanded_nodes, elapsed_ms 비교"],
        ["가시 검증", "blocked/visited/path layer를 켜고 실패 PoC 주변과 마지막 경로를 확인"],
    ], [2.2, 5.2])


def add_appendix(doc: Document) -> None:
    add_heading(doc, "6. C API와 C# Wrapper 대응표", 1)
    add_table(doc, ["C API", "C# wrapper"], [
        ["r3d_version", "Routing3DEngine.NativeVersion"],
        ["r3d_route_scene_text", "Routing3DEngine.RouteSceneText"],
        ["r3d_load_scene_text", "LoadSceneText"],
        ["r3d_set_grid / r3d_set_params", "SetGrid / SetParameters"],
        ["r3d_add_obstacle / r3d_add_passthrough", "AddObstacle / AddPassthrough"],
        ["r3d_add_task", "AddTask"],
        ["r3d_set_task_endpoints / diameter / goal_dir", "SetTaskEndpoints / SetTaskDiameter / SetTaskGoalDirection"],
        ["r3d_route_multi / route_ripup / route_task", "RouteMulti / RouteRipup / RouteTask"],
        ["r3d_route_corridor_multi / route_task_octree", "RouteCorridorMulti / RouteTaskOctree"],
        ["r3d_route_multi_progress", "RouteMultiProgress"],
        ["r3d_get_result / copy_path / copy_visited", "GetResult"],
        ["r3d_copy_blocked / copy_passthrough", "CopyBlocked / CopyPassthrough"],
        ["r3d_dump_scene_text / get_runtime_report", "DumpSceneText / GetRuntimeReportJson"],
        ["r3d_enum_octree_leaves", "EnumOctreeLeaves"],
    ], [3.4, 4.0])
    add_heading(doc, "6.1 실패 사유 enum", 2)
    add_table(doc, ["RouteFail", "의미"], [
        ["None", "성공 또는 실패 없음"],
        ["StartBlocked", "시작 셀이 막힘"],
        ["GoalBlocked", "목표 셀이 막힘"],
        ["CorridorMiss", "제한 corridor 안에 시작/목표 또는 연결 경로가 없음"],
        ["ExpansionLimit", "확장 노드 상한 도달"],
        ["GoalDirBlocked", "목표점에는 도달 가능하지만 요구 진입축을 만족하지 못함"],
        ["NoPath", "탐색 공간에서 연결 경로 없음"],
    ], [2.1, 5.3])


def main() -> None:
    doc = setup_doc()
    add_cover(doc)
    add_architecture(doc)
    add_quick_start(doc)
    add_c_api(doc)
    add_csharp_wrapper(doc)
    add_operational_guidance(doc)
    add_appendix(doc)
    add_para(doc, "출처 파일: cpp/capi/routing3d_capi.h, csharp/Routing3D.Engine/Native.cs, Routing3DEngine.cs, Models.cs, docs/routing3d_engine_reference.md", size=8.2, color="6B7280")
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
