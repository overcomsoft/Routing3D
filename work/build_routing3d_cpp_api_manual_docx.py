from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "cpp" / "docs" / "Routing3D_CPP_Engine_API_Manual_20260702.docx"
CURRENT = ROOT / "cpp" / "docs" / "Routing3D_CPP_Engine_API_Manual.docx"

FONT_BODY = "Malgun Gothic"
FONT_CODE = "Consolas"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "4B5563"
TABLE_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
NOTE_FILL = "EAF4EC"


def _rfonts(run, font: str) -> None:
    run.font.name = font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:eastAsia"), FONT_BODY)


def font(run, size: float, bold: bool = False, color: str = INK, name: str = FONT_BODY) -> None:
    _rfonts(run, name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def table_geometry(table, widths: list[float]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(w * 1440)))
        tbl_grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_width(cell, widths[i])
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def p(doc: Document, text: str = "", size: float = 10.0, bold: bool = False,
      color: str = INK, before: float = 0, after: float = 6) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.25
    r = para.add_run(text)
    font(r, size=size, bold=bold, color=color)


def heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.keep_with_next = True
    para.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    para.paragraph_format.space_after = Pt(10 if level == 1 else 7)
    r = para.add_run(text)
    font(r, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else DARK_BLUE)


def code(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, LIGHT_FILL)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(text.strip("\n").splitlines()):
        if i:
            para.add_run("\n")
        r = para.add_run(line)
        font(r, size=8.0, color="111827", name=FONT_CODE)


def note(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, NOTE_FILL)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    r = para.add_run(title + "  ")
    font(r, 9.5, True, "20603D")
    r = para.add_run(text)
    font(r, 9.5, False, INK)


def tbl(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_geometry(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, TABLE_FILL)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        r = para.add_run(h)
        font(r, 8.5, True, DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            para = cells[i].paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            for j, line in enumerate(text.split("\n")):
                if j:
                    para.add_run("\n")
                is_code = "`" in line or "::" in line or "(" in line and ")" in line
                r = para.add_run(line.replace("`", ""))
                font(r, 8.2, False, INK, FONT_CODE if is_code else FONT_BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def setup() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(10)
    return doc


def cover(doc: Document) -> None:
    p(doc, "Routing3D C++ Engine API Manual", 22, True, BLUE, after=3)
    p(doc, "C++ Core Header API, C ABI, Runtime Diagnostics", 14, True, "374151", after=10)
    p(doc, "갱신일: 2026-07-02 KST", 10, False, MUTED)
    p(doc, "대상: cpp/include/routing3d/*.hpp, cpp/capi/routing3d_capi.h, routing3d_capi.dll", 9.5, False, MUTED)
    note(doc, "문서 범위", "현재 저장소의 C++ 엔진과 C ABI 기준입니다. 모든 물리 좌표와 길이는 mm 단위이고, 라우팅 경로는 Cell(i,j,k) 정수 좌표 배열로 표현됩니다.")
    tbl(doc, ["섹션", "내용"], [
        ["1. 아키텍처", "C++ core, occupancy backend, routing algorithm, C ABI facade"],
        ["2. Core Header API", "geometry, occupancy, cost, A*, multi-route, scene I/O, octree"],
        ["3. C ABI", "DLL 생명주기, 입력, 옵션 setter, 라우팅, 결과 복사, trace/report"],
        ["4. 빌드/검증", "CMake, ctest, Viewer DLL 배치, 성능/진단 체크리스트"],
    ], [1.35, 5.15])
    doc.add_section(WD_SECTION.NEW_PAGE)


def architecture(doc: Document) -> None:
    heading(doc, "1. 엔진 구조", 1)
    p(doc, "Routing3D C++ 엔진은 SceneDoc를 중심 데이터 모델로 두고, 점유맵 backend와 A* 계열 탐색기를 교체 가능한 템플릿 계층으로 연결합니다.")
    code(doc, """
Client / Viewer / CLI
  -> C ABI facade: routing3d_capi.dll
    -> SceneDoc: grid + params + obstacles + tasks + results
      -> Occupancy: Dense, Sparse, Implicit, Octree, optional OpenVDB/FCL
        -> Routing: A*, Weighted A*, Segment A*, corridor, rip-up, CBS-lite
          -> Result: AStarResult / MultiRouteResult / R3dResult
""")
    tbl(doc, ["계층", "주요 파일", "책임"], [
        ["Geometry", "geometry.hpp", "Cell, Vec3, AABB, 6-neighbor, world-cell 변환"],
        ["Data model", "route_task.hpp, scene_io.hpp", "RouteTask, SceneDoc, SceneResult, scene v3 I/O"],
        ["Occupancy", "occupancy.hpp, box_index.hpp, octree_occupancy.hpp, vdb_occupancy.hpp", "blocked query, add_box, copy, sparse/implicit/octree backend"],
        ["Cost/search", "cost.hpp, astar.hpp", "RouteParams, clearance/tier/corridor cost, A*, Segment A*"],
        ["Multi-route", "multi_route.hpp, corridor.hpp", "priority ordering, mark_pipe, rip-up/CBS, hierarchical corridor"],
        ["Facade", "routing3d_capi.h/.cpp", "stable DLL ABI, UTF-8/POD/Safe interop contract"],
    ], [1.15, 2.25, 3.10])


def core_api(doc: Document) -> None:
    heading(doc, "2. Core Header API", 1)
    heading(doc, "2.1 geometry.hpp", 2)
    tbl(doc, ["타입/함수", "설명"], [
        ["Cell", "격자 좌표 i,j,k. 경로와 방문 셀의 기본 단위"],
        ["Vec3", "월드 좌표/크기. mm 단위"],
        ["AABB", "장애물/통과 객체의 axis-aligned bounding box"],
        ["NEIGHBORS_6", "고정 6방향 순서: +x, -x, +y, -y, +z, -z"],
        ["grid_cell_to_world / grid_world_to_cell", "cell 중심 좌표 변환과 floor 기반 world->cell 변환"],
        ["grid_box_range", "AABB가 차지하는 cell range 산출"],
    ], [2.35, 4.15])
    heading(doc, "2.2 occupancy.hpp", 2)
    tbl(doc, ["Backend", "특징", "주 사용처"], [
        ["DenseOccupancy", "1 byte/cell 연속 배열. O(1) query, copy 가능", "작은/중간 grid, golden test"],
        ["SparseOccupancy", "blocked cell만 hash set 저장", "초대형 corridor, sparse scene"],
        ["ImplicitOccupancy", "AABB를 SpatialBoxIndex에 보관하고 on-demand query", "5M+ grid 자동 전환, 메모리 절감"],
        ["OctreeOccupancy", "adaptive voxel tree와 jump query", "octree preview, macro-guide 진단"],
        ["VdbOccupancy", "OpenVDB optional backend", "압축 voxel backend 실험"],
    ], [1.55, 3.0, 1.95])
    heading(doc, "2.3 cost.hpp / astar.hpp", 2)
    tbl(doc, ["API", "역할"], [
        ["RouteParams", "cell_mm, w_turn, w_clear, w_heur, w_heur_near, clearance_radius, w_corridor, rack_levels, min_straight_cells, use_segment_astar"],
        ["RouteFail", "None, StartBlocked, GoalBlocked, CorridorMiss, ExpansionLimit, GoalDirBlocked, NoPath"],
        ["AStarResult", "success, path, length_mm, turns, expanded_nodes, cost_mm, elapsed_ms, fail, visited"],
        ["astar", "uniform-cost A*. 골든 검증용 최단 경로 탐색"],
        ["astar_weighted", "RouteParams 기반 weighted A*. turn/clearance/corridor/tier/goal_dir 지원"],
        ["astar_segmented", "straight-run expansion 기반 Segment A*. opt-in, cell-by-cell path 재구성"],
    ], [2.2, 4.3])
    heading(doc, "2.4 multi_route.hpp", 2)
    tbl(doc, ["API", "설명"], [
        ["order_indices / order_tasks", "priority: longest, shortest, diameter, utility, original"],
        ["snap_to_free_cell", "blocked start/goal을 근접 free cell로 보정"],
        ["mark_pipe", "성공 경로와 radius를 occupancy에 반영해 후속 배관 충돌 방지"],
        ["route_sequential", "작업 정렬 후 순차 라우팅. 성공 경로를 즉시 점유 처리"],
        ["route_ripup", "실패 route의 blocker를 제한적으로 걷어내고 재시도"],
        ["MultiRouteResult", "pipes, final occupancy, success_count, fail_count, total_length_mm, success_rate"],
    ], [2.2, 4.3])
    heading(doc, "2.5 scene_io.hpp / octree_occupancy.hpp", 2)
    tbl(doc, ["API", "설명"], [
        ["SCENE_FORMAT_VERSION = 3", "현재 scene format 버전"],
        ["Obstacle / SceneResult / SceneDoc", "AABB, task result, 전체 scene 데이터 모델"],
        ["dumps_scene / loads_scene", "SceneDoc 직렬화/역직렬화"],
        ["read_scene / write_scene", "UTF-8 scene 파일 I/O"],
        ["occupancy_from_doc / occupancy_from_passthrough", "SceneDoc에서 DenseOccupancy 생성"],
        ["OctreeOccupancy::build", "SceneDoc 장애물 기반 adaptive voxel map 생성"],
        ["astar_octree", "Octree leaf/jump 기반 single route 탐색"],
    ], [2.4, 4.1])


def c_abi(doc: Document) -> None:
    heading(doc, "3. C ABI: routing3d_capi.h", 1)
    note(doc, "ABI 원칙", "extern C, cdecl, x64, POD struct, UTF-8 string, opaque R3dEngine* handle을 사용합니다. C++ 예외는 ABI 경계를 넘지 않고 R3dStatus로 변환됩니다.")
    heading(doc, "3.1 Structs", 2)
    tbl(doc, ["Struct", "필드", "용도"], [
        ["R3dGrid", "cell_mm, origin, nx, ny, nz", "격자 정의"],
        ["R3dParams", "cost/heuristic/clearance/corridor/rack fields", "A* 비용과 탐색 옵션"],
        ["R3dRuntimeOptions", "large_grid_threshold, max_expansions, fallback_expansions, hier/ripup/cbs options", "대형 grid와 fallback 정책 제어"],
        ["R3dTraceOptions", "enabled, level, sample_every, include_occupancy, include_postprocess, max_events_per_task", "Trace JSONL 상세도 제어"],
        ["R3dResult", "success, length_mm, cost_mm, turns, expanded_nodes, elapsed_ms, path_len, visited_len, fail_reason", "task별 결과"],
        ["R3dOctreeLeaf", "x0/y0/z0, size_mm, state", "octree preview/diagnostics"],
    ], [1.7, 3.0, 1.8])
    heading(doc, "3.2 Lifecycle / setup", 2)
    tbl(doc, ["API", "설명"], [
        ["r3d_version / r3d_free_string", "버전 문자열 조회, caller-owned UTF-8 문자열 해제"],
        ["r3d_create / r3d_destroy", "엔진 핸들 생성/해제"],
        ["r3d_load_scene_text", "scene text를 엔진 상태로 적재"],
        ["r3d_set_grid / r3d_set_params", "격자와 비용 파라미터 설정"],
        ["r3d_set_runtime_options", "max expansion, large grid threshold, fallback/hier/CBS option 설정"],
        ["r3d_set_trace_options / r3d_set_trace_file / r3d_flush_trace", "Trace JSONL 설정과 flush"],
        ["r3d_get_runtime_report", "build flags, scene counts, effective options를 JSON으로 반환"],
    ], [2.8, 3.7])
    heading(doc, "3.3 Geometry / task input", 2)
    tbl(doc, ["API", "설명"], [
        ["r3d_add_obstacle", "AABB 장애물 추가. 라우팅 collision 대상"],
        ["r3d_add_passthrough", "통과 객체 추가. visualization/diagnostic용, collision 제외"],
        ["r3d_add_task", "start/end mm 좌표와 utility/group을 입력하고 task index 반환"],
        ["r3d_set_task_endpoints", "기존 task의 start/end 갱신"],
        ["r3d_set_task_diameter", "관경 기반 radius, diameter/utility priority 정렬에 사용"],
        ["r3d_set_task_goal_dir", "목표점 진입축 0..5 또는 -1(any) 설정"],
    ], [2.7, 3.8])
    heading(doc, "3.4 Routing commands", 2)
    tbl(doc, ["API", "설명", "기본 성격"], [
        ["r3d_route_scene_text", "scene text 입력 -> routed scene text 출력", "Level 1 convenience"],
        ["r3d_route_multi", "전체 task 순차 라우팅", "기본 batch route"],
        ["r3d_route_multi_progress", "route_multi + progress/cancel callback", "GUI/장시간 작업"],
        ["r3d_route_task", "단일 task를 원본 장애물 기준으로 탐색", "reroute/diagnostic"],
        ["r3d_route_task_anytime", "weighted A*를 시간 budget 안에서 점진 개선", "빠른 최초 경로 + 개선"],
        ["r3d_route_ripup", "blocker 기반 rip-up & reroute", "혼잡 회복"],
        ["r3d_route_corridor / r3d_route_corridor_multi", "coarse guide + fine tube corridor 탐색", "초대형 grid fallback"],
        ["r3d_route_task_octree", "Octree jump route", "single route diagnostic"],
    ], [2.25, 3.05, 1.2])
    heading(doc, "3.5 Option setters and result copy", 2)
    tbl(doc, ["API", "설명"], [
        ["r3d_set_corridor_cells", "외부 corridor seed IJK 배열 주입"],
        ["r3d_set_collect_visited", "visited cell 수집 on/off"],
        ["r3d_set_pipe_radius / r3d_set_per_task_radius / r3d_set_pipe_gap", "배관 반경/관경별 반경/배관 간 이격"],
        ["r3d_set_cbs_depth", "CBS-lite depth [0,3]"],
        ["r3d_set_min_straight / r3d_set_min_straight_mm", "코너 주변 최소 직선 길이 제약"],
        ["r3d_set_segment_astar / r3d_set_octree_guide / r3d_set_route_split", "대형/특수 경로 실험 옵션. 기본 OFF"],
        ["r3d_get_result / r3d_copy_path / r3d_copy_visited", "task 결과와 IJK 배열 복사"],
        ["r3d_copy_blocked / r3d_copy_blocked_sampled / r3d_copy_passthrough", "점유/통과 map 가시화용 복사"],
        ["r3d_dump_scene_text / r3d_enum_octree_leaves", "현재 scene dump와 octree leaf 열거"],
    ], [2.9, 3.6])


def examples(doc: Document) -> None:
    heading(doc, "4. C++ 사용 예시", 1)
    heading(doc, "4.1 Header-only core route", 2)
    code(doc, """
#include "routing3d/scene_io.hpp"
#include "routing3d/multi_route.hpp"
using namespace routing3d;

SceneDoc doc = read_scene("input.scene.txt");
DenseOccupancy occ = occupancy_from_doc(doc);

RouteParams params = doc.params;
params.w_turn = 500.0;
params.w_clear = 10.0;
params.w_heur = 1.0;

auto routed = route_sequential(occ, doc.tasks, params, "diameter",
                               /*pipe_radius*/0, /*snap_to_free*/2,
                               /*max_expansions*/-1, /*collect_visited*/false);

for (const PipeResult& p : routed.pipes) {
    if (p.result.success)
        printf("order=%d len=%.0f turns=%d\\n",
               p.order_index, p.result.length_mm, p.result.turns);
}
""")
    heading(doc, "4.2 C ABI route with result copy", 2)
    code(doc, """
#include "routing3d_capi.h"
#include <vector>

R3dEngine* e = r3d_create();
R3dGrid g{50.0, 0, 0, 0, 200, 160, 40};
r3d_set_grid(e, &g);

R3dParams p{};
p.cell_mm = 50.0;
p.w_turn = 500.0;
p.w_clear = 10.0;
p.w_heur = 1.0;
p.clearance_radius = 2;
p.clearance_connectivity = 6;
r3d_set_params(e, &p);

r3d_add_obstacle(e, 2000, 2000, 0, 3000, 3000, 2000);
int task = r3d_add_task(e, 100, 100, 500, 9000, 7000, 500, "ACID", "Exhaust");

r3d_set_task_diameter(e, task, 150.0);
r3d_set_per_task_radius(e, 1);
r3d_set_pipe_gap(e, 60.0);

r3d_route_multi(e, "diameter");
R3dResult result{};
r3d_get_result(e, task, &result);

std::vector<int32_t> ijk(static_cast<size_t>(result.path_len) * 3);
r3d_copy_path(e, task, ijk.data(), result.path_len);
r3d_destroy(e);
""")


def operations(doc: Document) -> None:
    heading(doc, "5. 빌드, 검증, 운영 체크리스트", 1)
    tbl(doc, ["목적", "명령/확인"], [
        ["C++ Release build", "cmake --build cpp/build --config Release --target routing3d_capi"],
        ["Core regression", "ctest --test-dir cpp/build -C Release --output-on-failure"],
        ["빠른 ABI 회귀", "ctest --test-dir cpp/build -C Release -R \"capi|golden|implicit|octree\""],
        ["Viewer 반영", "cpp/build/Release/routing3d_capi.dll을 Viewer bin/x64/Release 폴더로 복사"],
        ["C# solution build", "dotnet build csharp/Routing3D.Viewer.sln -c Release"],
        ["DB 실데이터 재현", "Routing3D.Viewer.exe --dbroute <projectId> <cellMm> <scope> <outPath>"],
    ], [2.0, 4.5])
    heading(doc, "5.1 실패 사유와 1차 대응", 2)
    tbl(doc, ["RouteFail", "원인", "대응"], [
        ["StartBlocked", "시작 cell이 blocked/out-of-bounds", "PoC snap, blocked map, 장비 내부 시작점 보정"],
        ["GoalBlocked", "목표 cell이 blocked/out-of-bounds", "terminal/duct 진입면 재투영, goal_dir 완화"],
        ["CorridorMiss", "fine corridor tube 밖에 start/goal 또는 연결 경로", "corridor radius/factor 확대 또는 direct fallback"],
        ["ExpansionLimit", "max_expansions 도달", "cell 상향, runtime options 상향, weighted heuristic/corridor 검토"],
        ["GoalDirBlocked", "goal cell 접근은 가능하지만 진입축 불만족", "goal_dir=-1 fallback, terminal stub 방향 확인"],
        ["NoPath", "연결 공간 없음 또는 선행 배관 교착", "rip-up/CBS, priority 변경, gap/radius A/B"],
    ], [1.35, 2.55, 2.6])
    heading(doc, "5.2 성능 판단 기준", 2)
    tbl(doc, ["수단", "효과", "주의"], [
        ["cell_mm 조정", "노드 수를 3D로 줄여 가장 큰 속도 개선", "좁은 통로 손실 가능. 실데이터 A/B 필요"],
        ["ImplicitOccupancy", "대형 grid에서 전체 voxel 배열 할당 회피", "AABB index query 비용은 장애물 분포에 의존"],
        ["collect_visited off", "메모리와 copy 비용 감소", "방문맵/애니메이션 사용 불가"],
        ["w_heur > 1", "확장 노드 감소", "최단성보다 속도 우선"],
        ["Segment/Octree/RouteSplit", "특정 구조에서 탐색 폭 축소 가능", "현재 기본 OFF. 실데이터 우월성 검증 후 적용"],
        ["CBS/rip-up", "혼잡 교착 회복", "분기 증가로 느려질 수 있어 depth/cap 제한 필요"],
    ], [1.8, 2.6, 2.1])


def build() -> None:
    doc = setup()
    cover(doc)
    architecture(doc)
    core_api(doc)
    c_abi(doc)
    examples(doc)
    operations(doc)
    p(doc, "주요 소스: cpp/capi/routing3d_capi.h, cpp/include/routing3d/{geometry,occupancy,cost,astar,multi_route,scene_io,octree_occupancy}.hpp", 8.2, False, MUTED)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    shutil.copyfile(OUT, CURRENT)
    print(OUT)
    print(CURRENT)


if __name__ == "__main__":
    build()
