from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\DINNO\DEV\AI-AutoRouting\Routing3D")
SRC = ROOT / "cpp" / "docs" / "Trace_Replay_Viewer_Usage.md"
OUT = ROOT / "cpp" / "docs" / "Trace_Replay_Viewer_Usage.docx"
SCREENSHOT = Path(r"C:\Users\overcom\AppData\Local\Temp\codex-clipboard-20f6c3c5-2197-47e6-9159-29f55a8810e4.png")

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)
TABLE_HEADER = "E8EEF5"
CODE_FILL = "F4F6F8"
CODE_BORDER = "C9D2DC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="AAB7C4", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, widths_in: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_in):
                cell.width = Inches(widths_in[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def apply_run_markup(paragraph, text: str, base_style: dict | None = None) -> None:
    base_style = base_style or {}
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(base_style.get("size", 10.5))
            run.font.color.rgb = RGBColor(0x9B, 0x1C, 0x1C)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = base_style.get("color", INK)
        else:
            run = paragraph.add_run(part)
            run.font.color.rgb = base_style.get("color", INK)


def add_code_block(doc: Document, code: str, lang: str = "") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, CODE_BORDER, "4")
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code.strip("\n"))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.8 if len(code.splitlines()) > 8 else 9.5)
    run.font.color.rgb = RGBColor(0x11, 0x27, 0x3A)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    set_table_borders(table)
    if cols == 2:
        widths = [1.85, 4.65]
    elif cols == 3:
        widths = [1.55, 1.75, 3.20]
    else:
        widths = [6.5 / cols] * cols
    set_table_width(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            apply_run_markup(p, text, {"size": 9.5})
            for run in p.runs:
                run.font.size = Pt(9.2)
            if r_idx == 0:
                set_cell_shading(cell, TABLE_HEADER)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = DARK_BLUE
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Routing3D Trace Replay Manual")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def parse_blocks(md: str):
    blocks = []
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("```"):
            lang = line.strip("`").strip()
            i += 1
            buf = []
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", lang, "\n".join(buf)))
            continue
        if line.lstrip().startswith("|"):
            buf = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                buf.append(lines[i])
                i += 1
            blocks.append(("table", buf))
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            blocks.append(("heading", level, text))
            i += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            buf = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i]):
                buf.append(re.sub(r"^\d+\.\s+", "", lines[i]).strip())
                i += 1
            blocks.append(("numbered", buf))
            continue
        if line.startswith("- "):
            buf = []
            while i < len(lines) and lines[i].startswith("- "):
                buf.append(lines[i][2:].strip())
                i += 1
            blocks.append(("bullets", buf))
            continue
        buf = [line.strip()]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith(("#", "```", "- ")) and not re.match(r"^\d+\.\s+", lines[i]) and not lines[i].lstrip().startswith("|"):
            buf.append(lines[i].strip())
            i += 1
        blocks.append(("paragraph", " ".join(buf)))
    return blocks


def build() -> None:
    md = SRC.read_text(encoding="utf-8")
    doc = Document()
    setup_styles(doc)
    add_footer(doc)

    title_done = False
    for block in parse_blocks(md):
        kind = block[0]
        if kind == "heading":
            _, level, text = block
            if level == 1 and not title_done:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(text)
                run.font.name = "Calibri"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
                run.font.size = Pt(24)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
                title_done = True
                subtitle = doc.add_paragraph("Routing3D.Viewer / Routing3D.TraceReplay 진단 창 사용 및 결과 해석")
                subtitle.paragraph_format.space_after = Pt(12)
                subtitle.runs[0].font.size = Pt(11)
                subtitle.runs[0].font.color.rgb = MUTED
                if SCREENSHOT.exists():
                    doc.add_picture(str(SCREENSHOT), width=Inches(6.5))
                    cap = doc.add_paragraph("그림 1. Routing3D Search Trace Replay 창 예시")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.runs[0].font.size = Pt(9)
                    cap.runs[0].font.color.rgb = MUTED
                continue
            style = "Heading 1" if level == 2 else "Heading 2" if level == 3 else "Heading 3"
            doc.add_paragraph(text, style=style)
        elif kind == "paragraph":
            p = doc.add_paragraph()
            apply_run_markup(p, block[1])
        elif kind == "bullets":
            for item in block[1]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.375)
                p.paragraph_format.first_line_indent = Inches(-0.188)
                p.paragraph_format.space_after = Pt(4)
                apply_run_markup(p, item)
        elif kind == "numbered":
            for item in block[1]:
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.left_indent = Inches(0.375)
                p.paragraph_format.first_line_indent = Inches(-0.188)
                p.paragraph_format.space_after = Pt(4)
                apply_run_markup(p, item)
        elif kind == "table":
            add_markdown_table(doc, block[1])
        elif kind == "code":
            add_code_block(doc, block[2], block[1])

    doc.save(OUT)


if __name__ == "__main__":
    build()
