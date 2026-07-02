from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "routing3d_route_modes_process_algorithm_expected_results_20260702.md"
OUT = ROOT / "docs" / "Routing3D_Route_Modes_Process_Algorithm_Expected_Results_20260702.docx"

FONT_BODY = "Malgun Gothic"
FONT_CODE = "Consolas"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "4B5563"
TABLE_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
NOTE_FILL = "EAF4EC"
CODE_FILL = "F3F5F7"


def rfonts(run, font_name: str = FONT_BODY) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), FONT_BODY)


def set_run(run, size: float = 11, bold: bool = False, color: str = INK, font_name: str = FONT_BODY) -> None:
    rfonts(run, font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def clean_inline(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    return text


def add_para(doc: Document, text: str, size: float = 11, bold: bool = False,
             color: str = INK, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(clean_inline(text))
    set_run(r, size=size, bold=bold, color=color)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(18 if level == 1 else 14 if level == 2 else 10)
    p.paragraph_format.space_after = Pt(10 if level == 1 else 7 if level == 2 else 5)
    r = p.add_run(clean_inline(text))
    set_run(r, size=16 if level == 1 else 13 if level == 2 else 12,
            bold=True, color=BLUE if level <= 2 else DARK_BLUE)


def add_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    set_run(r, size=20, bold=True, color="0B2545")


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(text)
    set_run(r, size=10.5, color=MUTED)


def add_callout(doc: Document, title: str, text: str, fill: str = NOTE_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + "  ")
    set_run(r, size=10, bold=True, color=DARK_BLUE)
    r = p.add_run(clean_inline(text))
    set_run(r, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc: Document, code: str, label: str | None = None) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, CODE_FILL)
    margins(cell, top=100, bottom=100, start=140, end=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if label:
        r = p.add_run(label + "\n")
        set_run(r, size=9, bold=True, color=DARK_BLUE)
    for i, line in enumerate(code.strip("\n").splitlines()):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        set_run(r, size=8.2, color="111827", font_name=FONT_CODE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [clean_inline(c.strip()) for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            continue
        rows.append(cells)
    return rows[0], rows[1:]


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not headers:
        return
    col_count = len(headers)
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    width = 6.5 / col_count
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, TABLE_FILL)
        margins(cell)
        set_cell_width(cell, width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run(r, size=8.8, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i in range(col_count):
            text = row[i] if i < len(row) else ""
            cell = cells[i]
            margins(cell)
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(text)
            set_run(r, size=8.6, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ["List Bullet", "List Number"]:
        s = styles[style_name]
        s.font.name = FONT_BODY
        s.font.size = Pt(11)
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing = 1.25
        s.paragraph_format.left_indent = Inches(0.375)
        s.paragraph_format.first_line_indent = Inches(-0.188)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Routing3D Route Modes | 2026-07-02")
    set_run(run, size=8.5, color=MUTED)


def add_lead_summary(doc: Document) -> None:
    add_callout(
        doc,
        "핵심 요약",
        "최단경로는 진단용 기준선, 기존설계추종은 기존 결과 재현용, 특징점반영은 실무 자동설계 기본 모드입니다. "
        "WTNHJ02 Exhaust 같은 실제 배관에서는 스텁, 접속면, 랙/회랑 패턴이 경로 품질에 큰 영향을 줍니다.",
        fill="EAF4EC",
    )


def build_doc() -> None:
    md = SRC.read_text(encoding="utf-8")
    lines = md.splitlines()

    doc = Document()
    set_document_defaults(doc)

    add_title(doc, "Routing3D 경로탐색 방식 3종 상세 문서")
    add_subtitle(doc, "최단경로 · 기존설계추종 · 특징점반영 | 프로세스, 알고리즘, 예상 결과")
    add_lead_summary(doc)

    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line.strip("`").strip()
                code_lines = []
            else:
                label = "프로세스 흐름" if code_lang == "mermaid" else "알고리즘/의사코드"
                add_code(doc, "\n".join(code_lines), label=label)
                in_code = False
                code_lang = ""
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip() or line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            # Title already rendered with a polished title block.
            i += 1
            continue

        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 1)
            i += 1
            continue

        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
            i += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            headers, rows = parse_table(table_lines)
            add_table(doc, headers, rows)
            continue

        bullet = re.match(r"^(\s*)-\s+(.*)$", line)
        if bullet:
            add_para(doc, bullet.group(2), style="List Bullet")
            i += 1
            continue

        numbered = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if numbered:
            add_para(doc, numbered.group(1), style="List Number")
            i += 1
            continue

        if line.endswith(":") and len(line) < 80:
            add_para(doc, line, size=10.5, bold=True, color=DARK_BLUE)
        else:
            add_para(doc, line)
        i += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Routing3D 경로탐색 방식 3종 상세 문서"
    doc.core_properties.subject = "Routing3D route mode process, algorithm, expected results"
    doc.core_properties.author = "DINNO / Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
