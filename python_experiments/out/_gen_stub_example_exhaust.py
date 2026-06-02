# -*- coding: utf-8 -*-
r"""기존설계 스텁 — 실제 Exhaust 배관 샘플 워크드 예시 문서(.docx) 생성기.

[실행]  (프로젝트 루트에서)
  .\.venv\Scripts\python.exe python_experiments/out/_gen_stub_example_exhaust.py
  powershell -ExecutionPolicy Bypass -File python_experiments/out/_docx_to_pdf.ps1 `
    -in docs/routing3d_stub_example_exhaust.docx -out docs/routing3d_stub_example_exhaust.pdf

산출물: docs/routing3d_stub_example_exhaust.docx

[문서 목적]
  DB(AUTOROUTINGV7, source_file=CLEAN_WTNHJ03_..._total.json)의 실제 Clean 장비 'WTNHJ02_'에서
  나가는 Exhaust(ACID) 배관 1개(GUID 2014e40a…, 150mm)를 표본으로, 출발(EQUIP)·종단(DUCT) 스텁이
  어떻게 추출·정규화·특징벡터화되고, 학습 대표값(집계 템플릿)이 어떻게 새 자동설계에 적용되는지를
  '실측 수치'로 보여준다. 본 문서의 모든 좌표/수치는 실제 DB 데이터를 학습 파이프라인
  (routing3d_py.pattern_learn.learn_pipe)에 통과시켜 캡처한 값이다(단위 mm).
  추출 명령: python -m routing3d_py.pattern_learn --project 6 --report
"""
import os

import matplotlib
matplotlib.use("Agg")  # 헤드리스 PNG 렌더(디스플레이 불필요).
import matplotlib.pyplot as plt
from mpl_toolkits.mplot3d.art3d import Line3DCollection  # noqa: F401 (3D 보장)

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "docs")

# ---- 실측 데이터(DB → learn_pipe 캡처). 본 표본 배관 GUID 2014e40a…, Exhaust/ACID, 150mm. ----
POLYLINE = [
    (187850, 9131, 15495), (187850, 9131, 15009), (187850, 9051, 14814),
    (187850, 9012, 14776), (187850, 8931, 14581), (187850, 8931, 13528),
    (188005, 8931, 13373), (190020, 8931, 13373), (190175, 8931, 13218),
    (190175, 8931, 13155), (190175, 8931, 12946), (190175, 8931, 12968),
    (190175, 8931, 12948),
]
EQUIP_AABB = ((185821, 5686, 15495), (190427, 16358, 17500))   # 장비 WTNHJ02_
DUCT_AABB = ((189920, 5008, 12448), (190720, 14493, 12948))     # LATERAL PIPE_db90d44a…
START_POC = POLYLINE[0]    # (187850, 9131, 15495) 출발(EQUIP, -z)
END_POC = POLYLINE[-1]     # (190175, 8931, 12948) 종단(DUCT, +z)
STUB_SPLIT_START = 6       # 점 0..6 = 출발 스텁(하강), 6..8 = 중간(수평), 8..12 = 종단 스텁
STUB_SPLIT_END = 8
CODE_FONT = "Consolas"
BODY_FONT = "Malgun Gothic"
CODE_BG = "F2F2F2"
HEAD_BG = "D9E2F3"


def _shade(elem_pr, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    elem_pr.append(shd)


def _set_run_font(run, ascii_font, size_pt, ea_font=None, bold=False, color=None):
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), ea_font or ascii_font)


def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rfonts.set(qn("w:eastAsia"), BODY_FONT)


def add_heading(doc, text, level):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    sizes = {0: 19, 1: 14.5, 2: 12, 3: 11}
    _set_run_font(run, BODY_FONT, sizes.get(level, 11), ea_font=BODY_FONT, bold=True,
                  color=(0x1F, 0x38, 0x64) if level <= 1 else (0x2E, 0x54, 0x96))
    return h


def add_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, BODY_FONT, 10.5, ea_font=BODY_FONT)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        _set_run_font(run, BODY_FONT, 10.5, ea_font=BODY_FONT)


def add_numbers(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(it)
        _set_run_font(run, BODY_FONT, 10.5, ea_font=BODY_FONT)


def add_code(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.0
    _shade(p._p.get_or_add_pPr(), CODE_BG)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line if line else "")
        _set_run_font(run, CODE_FONT, 8.6, ea_font=BODY_FONT)
        if i < len(lines) - 1:
            run.add_break()
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for j, htext in enumerate(headers):
        _shade(hdr[j]._tc.get_or_add_tcPr(), HEAD_BG)
        hdr[j].paragraphs[0].clear()
        run = hdr[j].paragraphs[0].add_run(htext)
        _set_run_font(run, BODY_FONT, 9.5, ea_font=BODY_FONT, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].paragraphs[0].clear()
            run = cells[j].paragraphs[0].add_run(str(val))
            _set_run_font(run, BODY_FONT, 9.0, ea_font=BODY_FONT)
    return t


def _box_edges(lo, hi):
    """AABB 12개 모서리 선분 [(p0,p1), …] 을 반환(와이어프레임용)."""
    x0, y0, z0 = lo
    x1, y1, z1 = hi
    c = [(x0, y0, z0), (x1, y0, z0), (x1, y1, z0), (x0, y1, z0),
         (x0, y0, z1), (x1, y0, z1), (x1, y1, z1), (x0, y1, z1)]
    e = [(0, 1), (1, 2), (2, 3), (3, 0), (4, 5), (5, 6), (6, 7), (7, 4),
         (0, 4), (1, 5), (2, 6), (3, 7)]
    return [(c[a], c[b]) for a, b in e]


def make_figure(png_path):
    """실제 표본 배관의 3D 다이어그램 + X-Z 측면 투영(2 패널)을 PNG 로 저장한다."""
    fig = plt.figure(figsize=(11, 4.6), dpi=150)

    # ---- (좌) 3D 뷰 ----
    ax = fig.add_subplot(1, 2, 1, projection="3d")
    for lo, hi, col, name in [(EQUIP_AABB[0], EQUIP_AABB[1], "#d9892f", "Equipment WTNHJ02_"),
                              (DUCT_AABB[0], DUCT_AABB[1], "#2f9d9d", "Lateral duct")]:
        for (p0, p1) in _box_edges(lo, hi):
            ax.plot(*zip(p0, p1), color=col, lw=0.8, alpha=0.6)
        ax.plot([], [], color=col, lw=2, label=name)
    # 스텁 구간별 색: 출발(빨강계)·중간(회색)·종단(파랑계).
    segs = [(0, STUB_SPLIT_START, "#e23030", "Start stub (EQUIP, -z)"),
            (STUB_SPLIT_START, STUB_SPLIT_END, "#888888", "Middle (free A*)"),
            (STUB_SPLIT_END, len(POLYLINE) - 1, "#3070ff", "End stub (DUCT, +z)")]
    for a, b, col, name in segs:
        xs = [POLYLINE[i][0] for i in range(a, b + 1)]
        ys = [POLYLINE[i][1] for i in range(a, b + 1)]
        zs = [POLYLINE[i][2] for i in range(a, b + 1)]
        ax.plot(xs, ys, zs, color=col, lw=2.4, label=name)
    ax.scatter(*START_POC, color="#e23030", s=55, depthshade=False)
    ax.scatter(*END_POC, color="#3070ff", s=55, depthshade=False)
    ax.set_title("3D — Exhaust(ACID) pipe + anchors", fontsize=9)
    ax.set_xlabel("X (mm)", fontsize=7); ax.set_ylabel("Y (mm)", fontsize=7)
    ax.set_zlabel("Z (mm)", fontsize=7)
    ax.tick_params(labelsize=6)
    ax.legend(fontsize=6, loc="upper left")
    ax.view_init(elev=18, azim=-58)

    # ---- (우) X-Z 측면 투영(하강→수평→상승 프로파일이 가장 또렷) ----
    ax2 = fig.add_subplot(1, 2, 2)
    for lo, hi, col, name in [(EQUIP_AABB[0], EQUIP_AABB[1], "#d9892f", "Equipment"),
                              (DUCT_AABB[0], DUCT_AABB[1], "#2f9d9d", "Duct")]:
        ax2.add_patch(plt.Rectangle((lo[0], lo[2]), hi[0] - lo[0], hi[2] - lo[2],
                                    fill=True, facecolor=col, alpha=0.12, edgecolor=col, lw=1.2))
    for a, b, col, _ in segs:
        xs = [POLYLINE[i][0] for i in range(a, b + 1)]
        zs = [POLYLINE[i][2] for i in range(a, b + 1)]
        ax2.plot(xs, zs, color=col, lw=2.6, marker="o", ms=3)
    ax2.scatter(START_POC[0], START_POC[2], color="#e23030", s=70, zorder=5)
    ax2.scatter(END_POC[0], END_POC[2], color="#3070ff", s=70, zorder=5)
    ax2.annotate("Start PoC (face -z)", (START_POC[0], START_POC[2]),
                 textcoords="offset points", xytext=(6, 6), fontsize=7, color="#e23030")
    ax2.annotate("End PoC (face +z)", (END_POC[0], END_POC[2]),
                 textcoords="offset points", xytext=(-30, -12), fontsize=7, color="#3070ff")
    ax2.set_title("X-Z side projection — down then up", fontsize=9)
    ax2.set_xlabel("X (mm)", fontsize=7); ax2.set_ylabel("Z (mm)", fontsize=7)
    ax2.tick_params(labelsize=6)
    ax2.grid(True, lw=0.3, alpha=0.4)

    fig.tight_layout()
    fig.savefig(png_path, bbox_inches="tight")
    plt.close(fig)
    print("figure:", png_path)


def render(title, subtitle, blocks, filename):
    doc = Document()
    set_base_style(doc)
    add_heading(doc, title, 0)
    sub = doc.add_paragraph()
    r = sub.add_run(subtitle)
    _set_run_font(r, BODY_FONT, 10.5, ea_font=BODY_FONT, color=(0x70, 0x70, 0x70))
    for kind, payload in blocks:
        if kind == "h1":
            add_heading(doc, payload, 1)
        elif kind == "h2":
            add_heading(doc, payload, 2)
        elif kind == "p":
            add_para(doc, payload)
        elif kind == "b":
            add_bullets(doc, payload)
        elif kind == "n":
            add_numbers(doc, payload)
        elif kind == "code":
            add_code(doc, payload)
        elif kind == "table":
            add_table(doc, payload[0], payload[1])
        elif kind == "img":
            doc.add_picture(payload[0], width=Inches(payload[1]))
            cap = doc.add_paragraph()
            cr = cap.add_run(payload[2])
            _set_run_font(cr, BODY_FONT, 9, ea_font=BODY_FONT, color=(0x70, 0x70, 0x70))
    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.normpath(os.path.join(OUT_DIR, filename))
    doc.save(path)
    print("saved:", path)


# =============================================================================
# 본문 — 실측 데이터(아래 값은 모두 DB → learn_pipe 캡처)
# =============================================================================
def build():
    os.makedirs(OUT_DIR, exist_ok=True)
    png = os.path.normpath(os.path.join(OUT_DIR, "routing3d_stub_example_exhaust_fig.png"))
    make_figure(png)
    blocks = [
        # ----------------------------------------------------------------- 1
        ("h1", "1. 샘플 — Clean 장비 Exhaust(ACID) 배관"),
        ("p", "DB(AUTOROUTINGV7)의 Clean 프로젝트(source_file=CLEAN_WTNHJ03_260417_035543_total.json, project_id=6)"
              "에서 실제 Exhaust 배관 한 개를 골라, 출발(EQUIP)·종단(DUCT) 스텁이 어떻게 만들어지고 자동설계에 "
              "쓰이는지 실측 수치로 따라간다. 이 프로젝트의 Exhaust 배관은 208개 중 20개, 그중 표본으로 다음을 쓴다."),
        ("table", [["항목", "값"], [
            ["ROUTE_PATH_GUID", "2014e40a-bde5-4985-b326-1f19d5a22534"],
            ["유틸리티 / 그룹", "ACID / Exhaust"],
            ["호칭경(외경)", "150 mm (SOURCE_SIZE→diameter_mm)"],
            ["출발 PoC(source_pos)", "(187850, 9131, 15495) — 메인 장비 WTNHJ02_"],
            ["종단 PoC(target_pos)", "(190175, 8931, 12948) — LATERAL PIPE_db90d44a…"],
            ["폴리라인 점 수", "13 (월드 mm, 순서대로)"],
        ]]),
        ("p", "사람이 그린 이 배관의 폴리라인(13점)은 다음과 같다. 장비 바닥(z=15495)에서 출발해 아래로 내려와 "
              "랙 높이(z≈13373)에서 수평으로 달린 뒤, 덕트 상부(z=12948)로 진입한다."),
        ("code",
         "  #   X        Y        Z          구간 해설\n"
         "   0 (187850,  9131, 15495)   ← 출발 PoC(장비 바닥면)\n"
         "   1 (187850,  9131, 15009)   │ 장비에서 아래로(-z)\n"
         "   2 (187850,  9051, 14814)   │\n"
         "   3 (187850,  9012, 14776)   │ 살짝 -y 로 비낌\n"
         "   4 (187850,  8931, 14581)   │\n"
         "   5 (187850,  8931, 13528)   │ 계속 하강 → 랙 진입\n"
         "   6 (188005,  8931, 13373)   ┐ 수평 런 시작(+x)\n"
         "   7 (190020,  8931, 13373)   ┘ 랙을 따라 +x 로 ~2000mm\n"
         "   8 (190175,  8931, 13218)   │ 덕트 위에서 하강\n"
         "   9 (190175,  8931, 13155)   │\n"
         "  10 (190175,  8931, 12946)   │\n"
         "  11 (190175,  8931, 12968)   │ (실측 미세 지터)\n"
         "  12 (190175,  8931, 12948)   ← 종단 PoC(덕트 상부면)"),
        ("img", (png, 6.6,
                 "그림 1. 실측 표본 배관의 3D 형상(좌)과 X-Z 측면 투영(우). 주황=장비 WTNHJ02_ AABB, "
                 "청록=레터럴 덕트 AABB. 빨강=출발 스텁(EQUIP, 장비 바닥 −z), 회색=중간 수평 런, "
                 "파랑=종단 스텁(DUCT, 덕트 상부 +z). 빨강 점=출발 PoC, 파랑 점=종단 PoC. "
                 "측면 투영에서 ‘아래로 하강 → 랙에서 수평 → 위로 진입’ 프로파일이 또렷하다.")),

        # ----------------------------------------------------------------- 2
        ("h1", "2. 출발(EQUIP) 스텁 추출"),
        ("p", "출발 PoC 를 포함/최근접하는 장비를 앵커로 잡는다. 여기선 메인 장비 WTNHJ02_ 의 AABB 안에 PoC 가 "
              "들어 있다. PoC 의 z(15495)가 장비 AABB 의 최소 z(15495)와 같아, PoC 는 장비 '바닥면(−z)'에 있다."),
        ("table", [["항목", "값", "의미"], [
            ["앵커(장비)", "WTNHJ02_", "find_equipment 가 PoC 포함 장비로 매칭"],
            ["anchor_min", "(185821, 5686, 15495)", "장비 AABB 하한"],
            ["anchor_max", "(190427, 16358, 17500)", "장비 AABB 상한"],
            ["poc_pos", "(187850, 9131, 15495)", "출발 PoC"],
            ["face", "−z", "PoC 가 가장 가까운 면 = 장비 바닥(도메인 규칙: 장비는 아래로 진출)"],
            ["dir_seq", "−z, −y, −z", "스텁 진행 방향열(축 스냅·연속 병합), 꺾임 2"],
            ["rise_mm", "1967", "면 법선축(z)으로 내려간 최대 거리(15495 → 13528)"],
            ["offset_mm", "0", "PoC 와 면 평면 간극(PoC 가 면 위에 있음)"],
        ]]),
        ("p", "추출 절차: 폴리라인을 출발 PoC 가 앞이 되도록 정렬한 뒤, PoC(점 0)에서 시작해 누적 4000mm 또는 "
              "꺾임 3회 한도까지 걸어간다(_walk_stub). 각 세그먼트를 가장 가까운 직교 축으로 스냅(axis_snap)하고 "
              "연속 동일 방향을 병합해 방향열을 만든다. 점 0→5 의 하강이 −z, 중간 −y 비낌, 다시 −z 로 정리돼 "
              "‘−z, −y, −z’ 가 된다. face 는 PoC 가 닿는 장비 면(−z)이다."),

        # ----------------------------------------------------------------- 3
        ("h1", "3. 종단(DUCT) 스텁 추출"),
        ("p", "종단 PoC 를 포함/최근접하는 덕트·레터럴을 앵커로 잡는다. 여기선 LATERAL PIPE_db90d44a… 의 AABB "
              "상단(z=12948)에 PoC 가 놓여, PoC 는 덕트 '상부면(+z)'에 있다. 종단 스텁은 종단 PoC 에서 출발 쪽으로 "
              "역순으로 걸어 추출한다."),
        ("table", [["항목", "값", "의미"], [
            ["앵커(덕트)", "LATERAL PIPE_db90d44a…", "find_duct 가 PoC 포함 덕트로 매칭"],
            ["anchor_min", "(189920, 5008, 12448)", "덕트 AABB 하한"],
            ["anchor_max", "(190720, 14493, 12948)", "덕트 AABB 상한"],
            ["poc_pos", "(190175, 8931, 12948)", "종단 PoC"],
            ["face", "+z", "PoC 가 가장 가까운 면 = 덕트 상부(도메인 규칙: 덕트는 위로 진입)"],
            ["dir_seq", "+z, −z, +z", "역순 진행 방향열(실측 지터로 미세 반전 포함), 꺾임 2"],
            ["rise_mm", "270", "면 법선축(z)으로 뜬 높이(덕트 상부 진입 구간)"],
            ["offset_mm", "0", "PoC 가 덕트 상부면 위에 있음"],
        ]]),
        ("p", "도메인 규칙이 수치로 드러난다: 출발은 face=−z(장비 아래로), 종단은 face=+z(덕트 위로). 같은 배관의 "
              "양 끝이 정확히 ‘아래로 나와 위로 들어간다’."),

        # ----------------------------------------------------------------- 4
        ("h1", "4. 특징벡터(24차원) — 두 스텁"),
        ("p", "각 스텁은 (앵커종류·유틸그룹·유틸) 범주 키 + 24차원 특징벡터로 표현된다. 구간은 "
              "[face 1hot 6][1차방향 6][2차방향 6][앵커내 상대좌표 3][시작→종단 단위 3]. 아래는 이 표본의 실측 벡터다."),
        ("code",
         "  EQUIP 스텁  key=(EQUIP, Exhaust, ACID)\n"
         "  feat = [0,0,0,0,0,1 | 0,0,0,0,0,1 | 0,0,0,1,0,0 | 0.441,0.323,0.000 | 0.673,-0.058,-0.737]\n"
         "          face=-z        dir1=-z       dir2=-y       rel(x,y,z)          dir_unit(s→t)\n"
         "  · rel z = 0.000  → PoC 가 장비 AABB 바닥(z 최소)에 있음을 수치로 확인\n"
         "\n"
         "  DUCT 스텁   key=(DUCT, Exhaust, ACID)\n"
         "  feat = [0,0,0,0,1,0 | 0,0,0,0,1,0 | 0,0,0,0,0,1 | 0.319,0.414,1.000 | -0.673,0.058,0.737]\n"
         "          face=+z        dir1=+z       dir2=-z       rel(x,y,z)          dir_unit(앵커접근=부호반전)\n"
         "  · rel z = 1.000  → PoC 가 덕트 AABB 상단(z 최대)에 있음을 수치로 확인"),
        ("p", "상대좌표 z 성분이 EQUIP=0.000(바닥)·DUCT=1.000(상단)으로 정확히 갈린다 — one-hot face 와 더불어 "
              "‘장비 아래/덕트 위’ 규칙을 벡터가 이중으로 담는다. dir_unit 은 출발→종단 단위벡터이며, 종단 스텁은 "
              "앵커로의 접근방향이라 부호가 반전된다(두 벡터가 서로 정확히 음수)."),

        # ----------------------------------------------------------------- 5
        ("h1", "5. 학습 대표값 — 집계 템플릿(n=17)"),
        ("p", "위 한 배관은 한 표본일 뿐이다. 같은 키의 Exhaust(ACID) 배관 17개를 모아 집계한 대표 스텁(뷰 "
              "route_stub_template)이 실제 추론에 쓰인다. 이 키는 면이 단일(EQUIP 17/17 −z, DUCT 17/17 +z)이라 "
              "대표 면이 명확하다."),
        ("table", [["키", "대표 face", "대표 dir_seq", "rise_mm", "avg_bends", "n"], [
            ["(EQUIP, Exhaust, ACID)", "−z", "−z, +y, −z", "1967", "1.8", "17"],
            ["(DUCT, Exhaust, ACID)", "+z", "+z, −z, +z", "297", "2.0", "17"],
        ]]),
        ("p", "면 분포가 100% 단일이므로, 이 키는 검색증강(L3b ANN)의 자기검증 게이트를 통과하지 못한다(분기할 "
              "이유가 없음) → 집계 대표면을 그대로 쓴다. 반면 UPW_S/HOT DI_S 처럼 면이 +x/+z 로 섞인 키만 ANN "
              "분기를 적용한다(본 표본 키는 해당 없음 = 깔끔한 단일면 예시)."),

        # ----------------------------------------------------------------- 6
        ("h1", "6. 자동설계 활용 — 새 Exhaust 배관 라우팅"),
        ("p", "이제 새 라우팅 작업(같은 키 Exhaust/ACID)의 PoC 에 학습 결과를 적용한다. 엔진(C++)은 그대로 두고, "
              "뷰어(C#)가 PoC 투영 면·회랑·랙을 결정해 넘긴다."),
        ("n", [
            "면 조회(L2a): PatternStore.TryGet(\"EQUIP\", \"Exhaust\", \"ACID\") → 면 −z, "
            "TryGet(\"DUCT\", \"Exhaust\", \"ACID\") → 면 +z.",
            "출발 PoC 투영: 장비 내부 PoC 를 학습 면(−z)으로, 즉 장비 바닥 바로 아래(+½셀 바깥)로 빼낸다"
            "(DropStartBelowEquipment + LiftPocToSurface(preferFace=\"-z\")). 고정 ‘최근접 면’이 아니라 학습 면을 쓴다.",
            "종단 PoC 투영: 덕트 솔리드 PoC 를 학습 면(+z), 즉 덕트 상부 바로 위로 투영"
            "(LiftPocToSurface(preferFace=\"+z\")) → 덕트 표면에 연결, 본체 관통 방지.",
            "접근불가 보정(전처리): 투영점이 다른 솔리드에 파묻히면 학습 면 방향으로 행진하며 최근접 자유셀로 "
            "스냅(SnapPocToFreeCell).",
            "랙 번들(L3a, 선택): Exhaust 그룹의 학습 랙 z-높이(이 프로젝트는 z≈13400 에 수평 런 97% 집중)를 "
            "엔진 rack_levels 로 주면, 이 배관도 같은 랙 높이에 수평 런이 모인다(위 표본의 z≈13373 런과 일치).",
            "회랑(L2b, 선택): 같은 출발·종단 PoC 의 기존 배관(이 GUID)을 찾으면 그 폴리라인을 회랑 셀로 주입해 "
            "새 경로가 사람 설계를 부드럽게 따라가게 한다.",
            "라우팅: 위로 투영·보정된 출발/종단점으로 가중 A*(w_heur=2.0)를 돌린다. 설비·덕트·레터럴·기설계 배관은 "
            "항상 솔리드 장애물이라 충돌을 회피한다.",
        ]),
        ("p", "결과적으로 새 Exhaust 배관은 학습대로 ‘장비 바닥(−z)에서 나와 → 랙 z-높이(≈13400)에서 수평으로 달려 "
              "→ 덕트 상부(+z)로 진입’하는, 위 표본과 같은 골격의 경로가 된다. 이 키는 단일면이라 면 예측이 100% "
              "정확하고, 랙 학습이 수평 런 높이를 사람 설계에 정렬한다."),
        ("code",
         "  새 PoC                학습 적용 후\n"
         "  start (장비 내부) ──▶ 장비 바닥 -z 바로 아래로 투영(학습 면)\n"
         "  goal  (덕트 솔리드)──▶ 덕트 상부 +z 바로 위로 투영(학습 면)\n"
         "  rack_levels = {z셀(13400 근처)}  → 수평 런이 공용 랙에 정렬\n"
         "  → 가중 A*(w_heur=2.0) → [장비 -z]──하강──[랙 +x 수평]──[덕트 +z]"),

        # ----------------------------------------------------------------- 7
        ("h1", "7. 재현 — CLI"),
        ("code",
         "  # 이 표본을 포함한 검수 리포트(키별 대표 면/방향/rise/꺾임)\n"
         "  python -m routing3d_py.pattern_learn --project 6 --report\n"
         "  # Exhaust 그룹 랙 z-높이(수평 런 집중) 리포트\n"
         "  python -m routing3d_py.pattern_learn --project 6 --rack-report\n"
         "  # pgvector 저장소 적재 후 키별 대표 템플릿 통계\n"
         "  python -m routing3d_py.pattern_learn --project 6 --write-db\n"
         "  python -m routing3d_py.pattern_db --stats\n"
         "  # 뷰어에서 Exhaust 그룹만 라우팅(헤드리스 A/B): 패턴/랙 on/off\n"
         "  Routing3D.Viewer.exe --dbroute 6 100 ACID out.txt   # R3D_PATTERNS/RACK/ANN env"),
        ("p", "주: 본 문서의 좌표·면·rise·특징벡터는 모두 위 --report 경로(DB → learn_pipe)로 캡처한 실측값이며, "
              "스텁 학습 메커니즘의 일반 설명은 별도 기술 레퍼런스(docs/routing3d_stub_pattern)를 참조한다."),
    ]
    render(
        "기존설계 스텁 — 실제 Exhaust(ACID) 배관 워크드 예시",
        "Routing3D · Clean 장비 WTNHJ02_ · GUID 2014e40a… · 150mm · 단위 mm (실측 데이터)",
        blocks,
        "routing3d_stub_example_exhaust.docx",
    )


if __name__ == "__main__":
    build()
