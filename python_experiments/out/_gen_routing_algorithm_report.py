# 라우팅 알고리즘 개발문서 docx 생성기 — docs/routing_algorithm_dev_report.md → .docx
# =============================================================================
# [실행]
#   .\.venv\Scripts\python.exe python_experiments/out/_gen_routing_algorithm_report.py
#   → docs/routing_algorithm_dev_report.docx 생성(Google Drive 에 올리면 Google Docs 로 열림)
#
# [하는 일]
#   마크다운(헤딩/표/코드펜스/불릿/인용/수평선/인라인 **굵게**·`코드`)을 파싱해
#   python-docx 로 서식 있는 .docx 를 만든다. 한글이 보이도록 본문/코드 East-Asian 폰트를
#   '맑은 고딕' 으로 지정한다(코드 ascii 는 Consolas).
# =============================================================================
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "docs" / "routing_algorithm_dev_report.md"
OUT = ROOT / "docs" / "routing_algorithm_dev_report.docx"

BODY_KR = "맑은 고딕"
CODE_ASCII = "Consolas"
CODE_BG = "F2F2F2"


def set_run_fonts(run, ascii_font, kr_font, size=None, color=None, bold=None):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), kr_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    if bold is not None:
        run.font.bold = bold


def shade_paragraph(p, fill):
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


# 인라인 **굵게** / `코드` 토큰화 → (text, bold, code) 리스트
INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_inline(p, text, base_ascii=BODY_KR, base_kr=BODY_KR, size=10.5):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2])
            set_run_fonts(r, base_ascii, base_kr, size=size, bold=True)
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1])
            set_run_fonts(r, CODE_ASCII, BODY_KR, size=size - 0.5,
                          color=(0xB0, 0x30, 0x60))
        else:
            r = p.add_run(tok)
            set_run_fonts(r, base_ascii, base_kr, size=size)


def add_code_block(doc, lines):
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Pt(6)
        shade_paragraph(p, CODE_BG)
        r = p.add_run(ln if ln else " ")
        set_run_fonts(r, CODE_ASCII, BODY_KR, size=9)


def add_table(doc, rows):
    # rows: 파싱된 셀 2차원(헤더 포함). 구분선(---) 행은 호출 전 제거.
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            txt = row[ci] if ci < len(row) else ""
            cell = cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline(p, txt, size=9)
            if ri == 0:
                for run in p.runs:
                    run.font.bold = True


def split_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def is_sep_row(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def main():
    md = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    # 기본 스타일 폰트(한글)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_KR
    normal.font.size = Pt(10.5)
    set_run_fonts(normal.element if False else doc.add_paragraph().add_run(""),
                  BODY_KR, BODY_KR)  # no-op 안전장치

    i = 0
    n = len(md)
    while i < n:
        line = md[i]
        stripped = line.rstrip()

        # 코드펜스
        if stripped.startswith("```"):
            j = i + 1
            buf = []
            while j < n and not md[j].rstrip().startswith("```"):
                buf.append(md[j])
                j += 1
            add_code_block(doc, buf)
            i = j + 1
            continue

        # 표(연속된 | 행)
        if stripped.startswith("|"):
            tbl = []
            j = i
            while j < n and md[j].lstrip().startswith("|"):
                if not is_sep_row(md[j]):
                    tbl.append(split_row(md[j]))
                j += 1
            if tbl:
                add_table(doc, tbl)
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
            i = j
            continue

        # 헤딩
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                h = doc.add_heading("", level=0)
            else:
                h = doc.add_heading("", level=min(level - 1, 4))
            add_inline(h, text, size=18 - level * 1.5)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x33, 0x55)
            i += 1
            continue

        # 수평선
        if re.match(r"^---+\s*$", stripped):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pbdr.append(bottom)
            pPr.append(pbdr)
            i += 1
            continue

        # 인용
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            add_inline(p, stripped.lstrip("> ").rstrip(), size=10)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # 불릿
        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^[-*]\s+", "", stripped), size=10.5)
            i += 1
            continue

        # 번호 목록
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", stripped), size=10.5)
            i += 1
            continue

        # 빈 줄
        if not stripped:
            i += 1
            continue

        # 일반 문단
        p = doc.add_paragraph()
        add_inline(p, stripped, size=10.5)
        i += 1

    doc.save(OUT)
    print(f"생성 완료: {OUT}  ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
