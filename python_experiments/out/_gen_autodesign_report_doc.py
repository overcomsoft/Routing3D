# -*- coding: utf-8 -*-
"""AI 자동설계 비교 리포트 — docx/xlsx 생성기 (C# 산출물 CSV+스냅샷을 공식 문서로).

[이 스크립트가 하는 일]
  C# 뷰어가 `--autodesign-report` 로 만든 출력 폴더(<group>_autodesign_report.csv +
  img/<group>_c###_{existing,shortest,stub}.png)를 읽어,
    · docx  — 전체 집계표 + 케이스별 지표표 + 3D 스냅샷 3-up(기존/최단/Stub+그룹)
    · xlsx  — 케이스×지표 시트(전략별 그룹핑F 최고값 강조) + 전체 집계 시트
  를 만든다. docx 는 선택적으로 MS Word COM 으로 PDF 변환(--pdf).
  서식 헬퍼는 _gen_spec_docs.py(set_base_style/add_heading/add_para/add_table) 재사용.

[선행]  먼저 C# 으로 데이터·스냅샷을 생성한다(DB 접속 필요):
  csharp/.../Routing3D.Viewer.exe --autodesign-report <projectId> <cellMm> <outDir>

[실행]  (프로젝트 루트에서)
  ./.venv/Scripts/python.exe python_experiments/out/_gen_autodesign_report_doc.py --in <outDir>
  ./.venv/Scripts/python.exe python_experiments/out/_gen_autodesign_report_doc.py --in <outDir> --pdf
  옵션: --out <docx경로>(기본 docs/routing3d_autodesign_report.docx) · --pdf(Word COM 변환)

[의존성]  python-docx, openpyxl (둘 다 .venv 설치됨).
"""
import argparse
import csv
import glob
import os
import subprocess
import sys

from docx import Document
from docx.shared import Inches, Pt

# 같은 폴더의 _gen_spec_docs 서식 헬퍼 재사용.
sys.path.insert(0, os.path.dirname(__file__))
from _gen_spec_docs import (  # noqa: E402
    set_base_style, add_heading, add_para, add_table, _set_run_font, BODY_FONT,
)

ROOT = os.path.normpath(os.path.join(os.path.dirname(__file__), "..", ".."))

# CSV 컬럼명(C# AutoDesignReport.WriteCsv 와 1:1) — 전략별 접두.
STRATS = [("기존", "existing"), ("최단", "shortest"), ("Stub그룹", "stub")]
STRAT_LABEL = {"기존": "기존설계", "최단": "최단(A*)", "Stub그룹": "Stub+그룹패턴"}


def _f(s):
    """문자열 셀 → float(빈/N/A 는 None)."""
    s = (s or "").strip().replace(",", "")
    if s in ("", "N/A"):
        return None
    try:
        return float(s)
    except ValueError:
        return None


def load_cases(out_dir):
    """출력 폴더의 *_autodesign_report.csv 를 읽어 (base, rows) 반환."""
    hits = glob.glob(os.path.join(out_dir, "*_autodesign_report.csv"))
    if not hits:
        raise SystemExit(f"CSV 없음: {out_dir}\\*_autodesign_report.csv (먼저 C# --autodesign-report 실행)")
    csv_path = hits[0]
    base = os.path.basename(csv_path)[: -len("_autodesign_report.csv")]
    with open(csv_path, encoding="utf-8-sig", newline="") as f:
        rows = list(csv.DictReader(f))
    return base, rows, csv_path


def img_path(out_dir, base, idx, suffix):
    """케이스 idx(1-based) 의 전략 스냅샷 절대경로(없으면 None)."""
    p = os.path.join(out_dir, "img", f"{base}_c{idx:03d}_{suffix}.png")
    return p if os.path.exists(p) else None


def fmt(v, nd=3):
    return "N/A" if v is None else f"{v:.{nd}f}"


def fmt0(v):
    return "N/A" if v is None else f"{v:,.0f}"


# =============================================================================
# docx
# =============================================================================
def build_docx(out_dir, base, rows, project_hint, out_docx):
    doc = Document()
    set_base_style(doc)
    add_heading(doc, "AI 자동설계 비교 리포트", 0)
    sub = doc.add_paragraph()
    r = sub.add_run(f"{project_hint} · 케이스 {len(rows)} · "
                    "기존설계 vs 자동설계(최단 / Stub+그룹패턴) — 길이·꺾임·그룹핑 Factor + 3D 뷰")
    _set_run_font(r, BODY_FONT, 10.5, ea_font=BODY_FONT, color=(0x70, 0x70, 0x70))

    # ---- 전체 집계 ----
    add_heading(doc, "1. 전체 집계", 1)
    agg = {p: {"ok": 0, "n": 0, "len": 0.0, "gf": []} for p, _ in STRATS}
    for row in rows:
        for p, _ in STRATS:
            if p == "기존":
                ok = int(row.get("작업수", "0") or 0)  # 기존은 매칭=작업수 근사(성공 컬럼 없음)
                n = ok
            else:
                succ = (row.get(f"{p}_성공", "0/0") or "0/0").split("/")
                ok, n = int(succ[0]), int(succ[1]) if len(succ) > 1 else 0
            agg[p]["ok"] += ok
            agg[p]["n"] += n
            agg[p]["len"] += _f(row.get(f"{p}_총길이mm")) or 0.0
            gf = _f(row.get(f"{p}_그룹핑F"))
            if gf is not None:
                agg[p]["gf"].append(gf)
    arows = []
    for p, _ in STRATS:
        a = agg[p]
        gfavg = sum(a["gf"]) / len(a["gf"]) if a["gf"] else None
        succ = f"{a['ok']}" if p == "기존" else f"{a['ok']}/{a['n']}"
        arows.append([STRAT_LABEL[p], succ, fmt0(a["len"]), fmt(gfavg)])
    add_table(doc, ["전략", "성공", "총길이(mm)", "그룹핑F(평균)"], arows)
    add_para(doc, "그룹핑 Factor = 0.35×랙집중도 + 0.30×번들밀집도 + 0.20×pitch일관성 + 0.15×레인정렬도 "
                  "(각 0~1, N/A 성분은 가중 재정규화). 최단=길이·직선성 우위, "
                  "Stub+그룹=그룹핑F(다발화)·사람설계 추종 우위가 기대값이다.")

    # ---- 케이스별 ----
    add_heading(doc, "2. 케이스별 비교", 1)
    cur_main = None
    for i, row in enumerate(rows, start=1):
        main = row.get("메인장비", "") or "(미상)"
        if main != cur_main:
            cur_main = main
            add_heading(doc, f"■ 메인장비: {main}", 2)
        equip = row.get("장비", "") or "(미상)"
        grp = row.get("유틸리티그룹", "") or "(미상)"
        ntask = row.get("작업수", "")
        add_heading(doc, f"케이스 {i}. {equip} / {grp} (작업 {ntask})", 3)

        crows = []
        for p, _ in STRATS:
            succ = "—" if p == "기존" else (row.get(f"{p}_성공", "") or "")
            crows.append([
                STRAT_LABEL[p], succ,
                fmt0(_f(row.get(f"{p}_총길이mm"))),
                fmt(_f(row.get(f"{p}_평균꺾임")), 1),
                fmt(_f(row.get(f"{p}_랙집중%")), 1),
                fmt(_f(row.get(f"{p}_번들밀집%")), 1),
                fmt(_f(row.get(f"{p}_pitch%")), 1),
                fmt(_f(row.get(f"{p}_lane%")), 1),
                fmt(_f(row.get(f"{p}_그룹핑F"))),
            ])
        add_table(doc, ["전략", "성공", "총길이(mm)", "평균꺾임", "랙집중%", "번들밀집%", "pitch%", "lane%", "그룹핑F"], crows)

        # 3D 스냅샷 3-up(있으면) — 1행 3열 표 셀에 그림 삽입.
        imgs = [(STRAT_LABEL[p], img_path(out_dir, base, i, suf)) for p, suf in STRATS]
        if any(path for _, path in imgs):
            t = doc.add_table(rows=2, cols=3)
            for c, (label, path) in enumerate(imgs):
                cell = t.rows[0].cells[c]
                cell.paragraphs[0].clear()
                if path:
                    cell.paragraphs[0].add_run().add_picture(path, width=Inches(2.05))
                cap = t.rows[1].cells[c]
                cap.paragraphs[0].clear()
                rr = cap.paragraphs[0].add_run(label)
                _set_run_font(rr, BODY_FONT, 8.5, ea_font=BODY_FONT)
        doc.add_paragraph()

    os.makedirs(os.path.dirname(out_docx), exist_ok=True)
    doc.save(out_docx)
    print("saved:", out_docx)
    return out_docx


# =============================================================================
# xlsx
# =============================================================================
def build_xlsx(base, rows, out_xlsx):
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "케이스별"
    head_fill = PatternFill("solid", fgColor="385B85")
    head_font = Font(bold=True, color="FFFFFF")
    best_fill = PatternFill("solid", fgColor="C6EFCE")   # 그룹핑F 최고 전략 강조(초록).

    header = ["메인장비", "장비", "유틸리티그룹", "작업수"]
    for p, _ in STRATS:
        lab = STRAT_LABEL[p]
        header += [f"{lab}\n성공", f"{lab}\n총길이mm", f"{lab}\n평균꺾임",
                   f"{lab}\n랙%", f"{lab}\n밀집%", f"{lab}\npitch%", f"{lab}\nlane%", f"{lab}\n그룹핑F"]
    ws.append(header)
    for c in range(1, len(header) + 1):
        cell = ws.cell(1, c)
        cell.fill = head_fill
        cell.font = head_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    gf_cols = {}   # 전략→그룹핑F 컬럼번호(1-based) — 행별 최고값 강조용.
    for ri, row in enumerate(rows, start=2):
        vals = [row.get("메인장비", ""), row.get("장비", ""), row.get("유틸리티그룹", ""),
                int(row.get("작업수", "0") or 0)]
        col = 5
        for p, _ in STRATS:
            succ = "—" if p == "기존" else (row.get(f"{p}_성공", "") or "")
            vals += [succ, _f(row.get(f"{p}_총길이mm")), _f(row.get(f"{p}_평균꺾임")),
                     _f(row.get(f"{p}_랙집중%")), _f(row.get(f"{p}_번들밀집%")),
                     _f(row.get(f"{p}_pitch%")), _f(row.get(f"{p}_lane%")), _f(row.get(f"{p}_그룹핑F"))]
            gf_cols[p] = col + 7   # 그룹핑F 는 전략 블록 8개 중 마지막.
            col += 8
        ws.append(vals)
        # 그룹핑F 최고 전략 강조.
        gfs = {p: _f(row.get(f"{p}_그룹핑F")) for p, _ in STRATS}
        best = max((v for v in gfs.values() if v is not None), default=None)
        if best is not None:
            for p, _ in STRATS:
                if gfs[p] is not None and abs(gfs[p] - best) < 1e-9:
                    ws.cell(ri, gf_cols[p]).fill = best_fill

    widths = [16, 14, 14, 7] + [9, 12, 9, 8, 8, 8, 8, 9] * len(STRATS)
    for c, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(c)].width = w
    ws.freeze_panes = "E2"

    # 전체 집계 시트.
    ws2 = wb.create_sheet("전체 집계")
    ws2.append(["전략", "성공", "총길이mm", "그룹핑F(평균)"])
    for c in range(1, 5):
        ws2.cell(1, c).fill = head_fill
        ws2.cell(1, c).font = head_font
    for p, _ in STRATS:
        oks = ns = 0
        lensum = 0.0
        gfs = []
        for row in rows:
            if p == "기존":
                k = int(row.get("작업수", "0") or 0)
                oks += k
                ns += k
            else:
                s = (row.get(f"{p}_성공", "0/0") or "0/0").split("/")
                oks += int(s[0])
                ns += int(s[1]) if len(s) > 1 else 0
            lensum += _f(row.get(f"{p}_총길이mm")) or 0.0
            g = _f(row.get(f"{p}_그룹핑F"))
            if g is not None:
                gfs.append(g)
        gfavg = sum(gfs) / len(gfs) if gfs else None
        ws2.append([STRAT_LABEL[p], f"{oks}" if p == "기존" else f"{oks}/{ns}",
                    round(lensum), round(gfavg, 3) if gfavg is not None else "N/A"])
    for c, w in enumerate([16, 10, 14, 13], start=1):
        ws2.column_dimensions[get_column_letter(c)].width = w

    os.makedirs(os.path.dirname(out_xlsx), exist_ok=True)
    wb.save(out_xlsx)
    print("saved:", out_xlsx)
    return out_xlsx


def to_pdf(docx_path):
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    ps1 = os.path.join(os.path.dirname(__file__), "_docx_to_pdf.ps1")
    try:
        subprocess.run(["powershell", "-ExecutionPolicy", "Bypass", "-File", ps1,
                        "-In", docx_path, "-Out", pdf_path], check=True, cwd=ROOT)
        print("saved:", pdf_path)
    except Exception as e:  # noqa: BLE001
        print("PDF 변환 실패(Word 필요):", e)


def main():
    ap = argparse.ArgumentParser(description="AI 자동설계 비교 리포트 docx/xlsx 생성")
    ap.add_argument("--in", dest="indir", required=True, help="C# --autodesign-report 출력 폴더")
    ap.add_argument("--out", dest="out",
                    default=os.path.join(ROOT, "docs", "routing3d_autodesign_report.docx"),
                    help="docx 출력 경로(기본 docs/routing3d_autodesign_report.docx)")
    ap.add_argument("--pdf", action="store_true", help="MS Word COM 으로 PDF 도 생성")
    args = ap.parse_args()

    out_dir = os.path.abspath(args.indir)
    base, rows, csv_path = load_cases(out_dir)
    project_hint = base
    print(f"읽음: {csv_path} (케이스 {len(rows)})")

    docx_path = os.path.abspath(args.out)
    build_docx(out_dir, base, rows, project_hint, docx_path)
    build_xlsx(base, rows, os.path.splitext(docx_path)[0] + ".xlsx")
    if args.pdf:
        to_pdf(docx_path)


if __name__ == "__main__":
    main()
