# -*- coding: utf-8 -*-
"""Step 1.1~1.4 상세 설계 문서(.docx) 생성기 — 구글독스 업로드용.

[실행]  (프로젝트 루트에서)
  .\.venv\Scripts\python.exe python_experiments/out/_gen_spec_docs.py

산출물: docs/spec/step1_1_occupancy.docx ... step1_4_multi_route.docx
구글독스: 파일 > 가져오기 > 업로드 하면 제목/표/목록/굵게가 거의 그대로 변환된다.
"""
import os

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "docs", "spec")

CODE_FONT = "Consolas"
BODY_FONT = "Malgun Gothic"
CODE_BG = "F2F2F2"
HEAD_BG = "D9E2F3"


# ----------------------------------------------------------------- low-level helpers

def _shade(elem_pr, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    elem_pr.append(shd)


def _set_run_font(run, ascii_font, size_pt, ea_font=None, bold=False, color=None):
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), ea_font or ascii_font)


def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rfonts.set(qn("w:eastAsia"), BODY_FONT)


def add_heading(doc, text, level):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    sizes = {0: 20, 1: 15, 2: 12.5, 3: 11}
    _set_run_font(run, BODY_FONT, sizes.get(level, 11), ea_font=BODY_FONT, bold=True,
                  color=(0x1F, 0x38, 0x64) if level <= 1 else (0x2E, 0x54, 0x96))
    return h


def add_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, BODY_FONT, 10.5, ea_font=BODY_FONT)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        _set_run_font(run, BODY_FONT, 10.5, ea_font=BODY_FONT)


def add_code(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.0
    _shade(p._p.get_or_add_pPr(), CODE_BG)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line if line else "")
        _set_run_font(run, CODE_FONT, 8.8, ea_font=BODY_FONT)
        if i < len(lines) - 1:
            run.add_break()
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for j, htext in enumerate(headers):
        _shade(hdr[j]._tc.get_or_add_tcPr(), HEAD_BG)
        hdr[j].paragraphs[0].clear()
        run = hdr[j].paragraphs[0].add_run(htext)
        _set_run_font(run, BODY_FONT, 9.5, ea_font=BODY_FONT, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].paragraphs[0].clear()
            run = cells[j].paragraphs[0].add_run(str(val))
            mono = j == 0 and len(headers) >= 2
            _set_run_font(run, CODE_FONT if mono else BODY_FONT, 9.0, ea_font=BODY_FONT)
    return t


def render(title, subtitle, blocks, filename):
    doc = Document()
    set_base_style(doc)
    add_heading(doc, title, 0)
    sub = doc.add_paragraph()
    r = sub.add_run(subtitle)
    _set_run_font(r, BODY_FONT, 10.5, ea_font=BODY_FONT, color=(0x70, 0x70, 0x70))
    for kind, payload in blocks:
        if kind == "h1":
            add_heading(doc, payload, 1)
        elif kind == "h2":
            add_heading(doc, payload, 2)
        elif kind == "h3":
            add_heading(doc, payload, 3)
        elif kind == "p":
            add_para(doc, payload)
        elif kind == "b":
            add_bullets(doc, payload)
        elif kind == "code":
            add_code(doc, payload)
        elif kind == "table":
            add_table(doc, payload[0], payload[1])
    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.normpath(os.path.join(OUT_DIR, filename))
    doc.save(path)
    print("saved:", path)


# =============================================================================
# Step 1.1 — 점유맵
# =============================================================================
def doc_occupancy():
    blocks = [
        ("h1", "1. 개요"),
        ("p", "플랜트 공간을 일정 크기(기본 50mm)의 정육면체 셀로 나눈 3차원 격자로 표현하고, "
               "각 셀이 '장애물로 점유되었는지'를 True/False 로 관리한다. 이후 직교 A* 탐색이 "
               "이 점유맵에 '이 셀로 지나갈 수 있는가?'를 질의한다."),
        ("p", "등간격 복셀은 고해상도/대영역에서 셀 수가 폭발한다. 그래서 같은 질의 인터페이스"
               "(OccupancyMap 추상 클래스) 뒤에 세 가지 저장 방식을 두고 상황에 맞게 고른다. "
               "A* 등 사용자 코드는 인터페이스(is_blocked / in_bounds / bounds / to_world / to_cell)"
               "에만 의존하므로 백엔드를 자유롭게 교체할 수 있다."),

        ("h1", "2. 단위·좌표 규약"),
        ("b", [
            "모든 월드 좌표·치수의 단위는 밀리미터(mm).",
            "셀 인덱스는 정수 3-튜플 (i, j, k). 셀 (0,0,0) 은 origin 에서 시작.",
            "셀 중심 월드좌표 = origin + (cell + 0.5) × cell_mm.",
            "shape = (nx, ny, nz) 는 각 축의 셀 개수.",
            "격자 범위 밖의 셀은 항상 '점유(True)'로 간주 → A* 가 격자 밖으로 못 나간다.",
        ]),

        ("h1", "3. 전체 구조 흐름도"),
        ("p", "추상 베이스가 좌표변환·복셀화 등 '저장 방식과 무관한' 로직을 공통 구현하고, "
               "실제 점유 비트의 저장/조회만 백엔드별로 구현한다."),
        ("code",
         "                 ┌──────────────────────────────────────────┐\n"
         "   사용자 코드     │   OccupancyMap (추상 베이스, ABC)         │\n"
         "   (A*, 비용함수,  │  공통: in_bounds / bounds / world_bounds  │\n"
         "    시각화)        │        to_world / to_cell / is_blocked    │\n"
         "      │            │        add_box / add_boxes / block_cell   │\n"
         "      └──질의────▶ │  추상: _init_storage / _get / _set        │\n"
         "                   │        count_blocked / inflate / copy     │\n"
         "                   │        approx_bytes / to_numpy            │\n"
         "                   └───────┬─────────────┬─────────────┬───────┘\n"
         "                           │ 상속        │ 상속        │ 상속\n"
         "                  ┌────────▼───┐ ┌───────▼──────┐ ┌────▼─────────┐\n"
         "                  │ Dense      │ │ BitPacked    │ │ Sparse       │\n"
         "                  │ np.bool[]  │ │ uint8 1비트  │ │ set[(i,j,k)] │\n"
         "                  │ 셀당 1byte │ │ 셀당 1bit    │ │ 점유만 저장  │\n"
         "                  └────────────┘ └──────────────┘ └──────────────┘"),

        ("h1", "4. 백엔드 비교"),
        ("table", (
            ["백엔드", "셀당 메모리", "질의 속도", "적합 상황", "실측(25.6% 점유, 50mm)"],
            [
                ["DenseOccupancyMap", "1 byte", "O(1) 최속", "작은 ROI, 최속 질의(기본값)", "625 KB"],
                ["BitPackedOccupancyMap", "1 bit", "약간 느림", "같은 메모리로 ~8배 큰 ROI", "78 KB"],
                ["SparseOccupancyMap", "set 엔트리", "해시 조회", "점유가 희박할 때만", "23.6 MB"],
            ],
        )),
        ("p", "주의: 바닥·기둥처럼 점유가 빽빽하면 Sparse 는 set 오버헤드 때문에 오히려 Dense 의 "
               "~37배까지 커진다. 균일한 큰 덩어리의 진짜 압축(옥트리/VDB)은 433m 스케일이 목표인 "
               "Phase 3(OpenVDB)에서 도입한다."),

        ("h1", "5. 핵심 알고리즘"),
        ("h2", "5.1 좌표 변환"),
        ("b", [
            "to_world(cell) = origin + (cell + 0.5) × cell_mm  → 셀 중심 월드좌표(mm).",
            "to_cell(world) = floor((world − origin) / cell_mm)  → 포함 셀 인덱스.",
        ]),
        ("h2", "5.2 add_box — AABB 복셀화"),
        ("code",
         "1) lo/hi 월드좌표 → 셀 범위 (시작 = floor, 끝 = ceil, 끝은 제외 경계)\n"
         "2) 격자 [0, shape) 로 클리핑\n"
         "3) 비어 있으면 0 반환, 아니면 _fill_box 로 채우고 '신규' 점유 셀 수 반환\n"
         "   - Dense: NumPy 슬라이스 대입으로 가속 (_fill_box 오버라이드)\n"
         "   - Sparse/BitPacked: 셀 단위 3중 루프(기본 구현)"),
        ("h2", "5.3 inflate — 장애물 팽창(하드 클리어런스)"),
        ("p", "장애물을 radius_cells 만큼 부풀린 같은 백엔드의 새 점유맵을 만든다. 1단계 팽창은 "
               "_dilate_once(이웃 방향으로 OR-시프트)로 수행하며 radius_cells 번 반복한다. "
               "connectivity 6 = 맨해튼 볼(면 인접), 26 = 체비셰프 볼(면+모서리+꼭짓점)."),
        ("code",
         "_shift(grid, di,dj,dk) : 그리드를 (di,dj,dk)만큼 평행이동(범위 밖은 0)\n"
         "_dilate_once(grid, off): result = grid; for o in off: result |= shift(grid,o)\n"
         "inflate(r)             : for _ in range(r): grid = _dilate_once(grid, off)"),

        ("h1", "6. 자료구조 / 클래스"),
        ("table", (
            ["이름", "종류", "설명"],
            [
                ["AABB", "frozen dataclass", "장애물 직육면체. lo<hi 검증. 필드 lo/hi = (x,y,z) mm."],
                ["OccupancyMap", "ABC", "질의 인터페이스 + 공통 로직(좌표/복셀화). 백엔드 공통 계약."],
                ["DenseOccupancyMap", "구현", "np.bool 3D 배열(grid). 가장 단순·최속."],
                ["BitPackedOccupancyMap", "구현", "uint8 배열(packed), z 8셀=1byte, little 비트순서."],
                ["SparseOccupancyMap", "구현", "점유 셀 좌표 set(blocked)."],
            ],
        )),

        ("h1", "7. 주요 함수"),
        ("table", (
            ["함수 / 메서드", "역할"],
            [
                ["__init__(shape, origin, cell_mm)", "shape/origin/cell_mm 검증·저장 후 백엔드 초기화."],
                ["from_world_bounds(lo, hi, cell_mm)", "월드 범위를 덮는 빈 맵 생성(셀 단위 올림)."],
                ["in_bounds(cell)", "셀이 격자 범위 안인지."],
                ["is_blocked(cell)", "점유 또는 격자 밖이면 True (A* 질의 핵심)."],
                ["block_cell(cell)", "단일 셀 점유 표시(격자 밖 무시)."],
                ["add_box(box) / add_boxes(boxes)", "AABB(들) 복셀화. 신규 점유 셀 수 반환."],
                ["to_world(cell) / to_cell(world)", "셀 ↔ 월드좌표(mm) 변환."],
                ["world_bounds() / bounds()", "격자가 덮는 월드 AABB / 셀 인덱스 범위."],
                ["count_blocked()", "점유 셀 총 개수(백엔드별 구현)."],
                ["inflate(radius, connectivity)", "팽창된 새 맵 반환(백엔드별 구현)."],
                ["to_numpy()", "(nx,ny,nz) bool 배열로 펼침(시각화/내보내기/거리변환)."],
                ["copy()", "독립 복사본(다중 배관 순차 라우팅에서 원본 보존용)."],
                ["approx_bytes()", "저장소 추정 바이트(백엔드 메모리 비교용)."],
            ],
        )),

        ("h1", "8. 주요 변수 / 상수"),
        ("table", (
            ["이름", "타입", "설명"],
            [
                ["shape", "(nx,ny,nz) int", "각 축 셀 개수."],
                ["origin", "np.float64[3]", "격자 원점 월드좌표(mm)."],
                ["cell_mm", "float", "셀 한 변 길이(mm). 기본 50, 설정 가능."],
                ["Cell", "타입별칭", "셀 인덱스 (i,j,k) 정수 3-튜플."],
                ["NEIGHBORS_6", "tuple[Cell]", "면 인접 6방향(±x,±y,±z). 직교 이동/6-연결 팽창."],
                ["NEIGHBORS_26", "tuple[Cell]", "면+모서리+꼭짓점 26방향. 26-연결 팽창."],
                ["grid / packed / blocked", "백엔드 저장소", "Dense/BitPacked/Sparse 의 점유 저장 필드."],
            ],
        )),

        ("h1", "9. 실행 명령어"),
        ("code",
         "# 단위 테스트 (python_experiments/ 에서)\n"
         "..\\.venv\\Scripts\\python.exe -m pytest tests/test_occupancy.py -v\n\n"
         "# 코드에서 사용\n"
         "from routing3d_py import DenseOccupancyMap, BitPackedOccupancyMap, AABB\n"
         "occ = DenseOccupancyMap((40, 40, 40), origin=(0,0,0), cell_mm=50)\n"
         "occ.add_box(AABB((100,100,100), (900,900,200)))   # 신규 점유 셀 수 반환\n"
         "occ.is_blocked((2, 2, 2))                          # A* 질의"),
    ]
    render("Phase 1 · Step 1.1 — 점유맵(Occupancy Map)",
           "Routing3D 플랜트 배관 3D 직교 라우팅 · 단위 mm · 기본 셀 50mm · 모듈 occupancy.py",
           blocks, "step1_1_occupancy.docx")


# =============================================================================
# Step 1.2 — 직교 A*
# =============================================================================
def doc_astar():
    blocks = [
        ("h1", "1. 개요"),
        ("p", "점유맵(OccupancyMap, 어느 백엔드든) 위에서 시작 셀 → 목표 셀까지의 최단 직교 경로를 "
               "A* 알고리즘으로 찾는다. 이동은 6방향(±X,±Y,±Z) 직교만 허용하며 대각선은 금지한다."),
        ("b", [
            "Step 1.2 범위: 비용은 균일 이동비용(셀 1칸 = cell_mm)만 사용.",
            "휴리스틱은 맨해튼 거리(셀 수) × cell_mm → 직교 격자에서 admissible & consistent.",
            "turn penalty·클리어런스·단 분리 등 비용 튜닝은 Step 1.3(cost.py)에서 확장한다.",
        ]),

        ("h1", "2. A* 탐색 흐름도"),
        ("code",
         "  open(우선순위 큐): f = g + h 가 작은 것부터 꺼냄\n"
         "        │\n"
         "        ▼  pop current\n"
         "  current == goal ? ──예──▶ came_from 역추적으로 경로 복원 → 반환\n"
         "        │ 아니오\n"
         "        ▼  6-이웃 nb 마다\n"
         "  is_blocked(nb)? ──예──▶ 건너뜀 (장애물 / 격자 밖)\n"
         "        │ 아니오\n"
         "        ▼  g_new = g[current] + step_cost\n"
         "  g_new < g[nb]? ──예──▶ g[nb]=g_new, came_from[nb]=current, open 에 push\n"
         "        │\n"
         "        └ open 이 비면 → 경로 없음(None)"),

        ("h1", "3. 핵심 알고리즘"),
        ("b", [
            "f(n) = g(n) + h(n). g = 시작부터 누적 실비용, h = 목표까지 추정(휴리스틱).",
            "admissibility: h ≤ 실제 최단비용 → A* 는 최적 경로를 보장. "
            "한 칸 비용이 정확히 cell_mm 이고 h = 맨해튼×cell_mm 이므로 과대평가가 없다.",
            "consistency: 인접 셀로 갈 때 h 감소량 ≤ 이동비용 → closed 재방문 불필요.",
            "tie-break: heapq 항목에 단조 증가 counter 를 넣어 f 동률 시 안정 정렬·비교 오류 방지.",
            "closed 집합으로 이미 확정된 셀의 낡은 힙 항목(중복 pop)을 무시한다.",
        ]),

        ("h1", "4. 자료구조"),
        ("table", (
            ["이름", "타입", "설명"],
            [
                ["open_heap", "heapq[(f,counter,cell)]", "f 최소 힙. counter 로 안정 tie-break."],
                ["g", "dict[cell→float]", "셀별 최소 누적 비용(mm). 미방문은 +∞ 취급."],
                ["came_from", "dict[cell→cell]", "경로 복원용 직전 셀 맵."],
                ["closed", "set[cell]", "확정(expand)된 셀. 중복 pop 무시."],
                ["AStarResult", "dataclass", "탐색 결과 묶음(아래 표)."],
            ],
        )),
        ("h2", "AStarResult 필드"),
        ("table", (
            ["필드", "설명"],
            [
                ["success", "경로를 찾았는지 여부."],
                ["path", "셀 리스트 [start..goal]. 실패 시 None."],
                ["length_mm", "경로 기하 길이(mm) = (셀 수 − 1) × cell_mm."],
                ["turns", "방향 전환(직각 회전) 횟수."],
                ["expanded_nodes", "확장한 노드(상태) 수 = 탐색 비용 지표."],
                ["visited", "확장된 모든 셀(collect_visited=True 일 때). 시각화용."],
                ["elapsed_ms", "탐색 소요 시간(ms)."],
                ["cost_mm", "총 비용(mm). 균일 A* 에서는 length_mm 과 동일."],
            ],
        )),

        ("h1", "5. 주요 함수"),
        ("table", (
            ["함수", "역할"],
            [
                ["astar(occ, start, goal, *, step_cost, collect_visited, max_expansions)",
                 "균일 비용 직교 A*. 상태 = 셀."],
                ["astar_weighted(occ, start, goal, params, ...)",
                 "비용함수 A*(Step 1.3). 상태 = (셀, 진입방향). cost.py 참조."],
                ["astar_world(occ, start_mm, goal_mm, **kw)",
                 "월드 좌표(mm)를 셀로 변환 후 astar 호출하는 편의 함수."],
                ["manhattan(a, b)", "두 셀의 맨해튼 거리(셀 수). 휴리스틱 기반."],
                ["count_turns(path)", "경로의 방향 전환 횟수 계산."],
                ["_reconstruct(came_from, goal)", "came_from 역추적 → 정방향 경로 리스트."],
            ],
        )),

        ("h1", "6. 주요 변수 / 파라미터"),
        ("table", (
            ["이름", "기본값", "설명"],
            [
                ["step_cost", "None→cell_mm", "셀 1칸 이동 비용(mm)."],
                ["collect_visited", "False", "True 면 확장 셀을 모두 모아 시각화/디버그에 사용."],
                ["max_expansions", "None", "확장 셀 수 상한(폭주 방지). None=무제한."],
                ["sc", "cell_mm", "내부: 실제 step_cost."],
                ["expanded", "0", "내부: 누적 확장 노드 수."],
            ],
        )),

        ("h1", "7. 실패(success=False) 조건"),
        ("b", [
            "start 또는 goal 이 점유 / 격자 밖.",
            "목표까지 경로가 없음(open 소진).",
            "max_expansions 초과.",
        ]),

        ("h1", "8. 실행 명령어"),
        ("code",
         "# DB 영역에서 start→goal(mm) 경로 탐색 + 지표 출력\n"
         ".\\.venv\\Scripts\\python.exe -m routing3d_py.astar ^\n"
         "    --region 195000 8000 14000 205000 12000 16000 --cell-mm 50 ^\n"
         "    --start 196000 9000 15600 --goal 204000 11000 15600\n\n"
         "# 경로+점유맵 3D 렌더(스크린샷)\n"
         ".\\.venv\\Scripts\\python.exe -m routing3d_py.astar ^\n"
         "    --region 195000 8000 14000 205000 12000 16000 ^\n"
         "    --start 196000 9000 15600 --goal 204000 11000 15600 ^\n"
         "    --screenshot python_experiments/out/route.png\n\n"
         "# 단위 테스트\n"
         ".\\.venv\\Scripts\\python.exe -m pytest python_experiments/tests/test_astar.py -v"),
    ]
    render("Phase 1 · Step 1.2 — 직교 A* 경로 탐색",
           "Routing3D · 6방향 직교 이동 · 맨해튼×cell_mm 휴리스틱 · 모듈 astar.py",
           blocks, "step1_2_astar.docx")


# =============================================================================
# Step 1.3 — 비용함수
# =============================================================================
def doc_cost():
    blocks = [
        ("h1", "1. 개요"),
        ("p", "직교 A* 의 '이동 비용'을 정의한다. 기본 이동비용(셀 = cell_mm) 위에 다음을 더한다."),
        ("b", [
            "Turn penalty: 진행 방향이 바뀔 때 가산 → 직각 회전 최소화(엘보 줄이기).",
            "클리어런스 페널티: 장애물에 가까운 셀일수록 가산 → 벽에서 떨어져 지나가게 유도.",
            "단(段) 분리: z 레벨별 가산 → 특정 단(배관 랙)으로 유도/회피.",
        ]),

        ("h1", "2. 핵심 설계 결정 — 왜 '보너스'가 아니라 '페널티'인가"),
        ("p", "계획서 초안은 클리어런스를 '여유 셀 수에 비례한 비용 감산(보너스)'으로 표현했다. "
               "그러나 비용을 감산하면 한 칸 이동 비용이 cell_mm 보다 작아질 수 있어, "
               "맨해튼×cell_mm 휴리스틱이 실제 비용을 과대평가하게 되고 → A* 의 admissibility(최적성)가 깨진다."),
        ("p", "그래서 동일한 목적(벽에서 멀어지기)을 '장애물 근접 시 가산하는 페널티'로 구현한다. "
               "모든 가산항이 ≥ 0 이므로 이동 비용 ≥ cell_mm 이 보장되고, 맨해튼 휴리스틱이 "
               "admissible & consistent 하게 유지된다(A* 최적성 보존)."),

        ("h1", "3. 전체 흐름도"),
        ("code",
         "  RouteParams(가중치)  +  OccupancyMap\n"
         "        │\n"
         "        ▼  CostModel(occ, params)        # 생성 시 클리어런스 맵 1회 사전계산\n"
         "        │\n"
         "        ▼  move_cost(to_cell, prev_off, move_off)\n"
         "        =  cell_mm\n"
         "           + (w_turn          if 방향 바뀜)\n"
         "           + 클리어런스 페널티(to_cell)    # 장애물에 가까울수록 큼\n"
         "           + 단 분리 페널티(to_cell.z)"),

        ("h1", "4. 핵심 알고리즘"),
        ("h2", "4.1 clearance_map — bounded distance transform"),
        ("p", "각 셀에서 '가장 가까운 장애물까지의 거리(셀)'를 max_radius 로 상한해 계산한다. "
               "장애물을 1단계씩 팽창하며 새로 덮이는 셀에 그 단계 번호를 거리로 기록한다."),
        ("code",
         "장애물 셀 = 거리 0,  멀리 떨어진(끝까지 안 덮인) 셀 = max_radius\n"
         "current = grid(장애물)\n"
         "for d in 1..max_radius:\n"
         "    dilated = _dilate_once(current, offsets)\n"
         "    newly   = dilated & ~current      # 이번 단계에 처음 덮인 셀\n"
         "    dist[newly] = d\n"
         "    current = dilated"),
        ("h2", "4.2 cell_penalty — 목적지 셀 가산 페널티"),
        ("code",
         "pen = 0\n"
         "if clearance 있음 and d < clearance_radius:\n"
         "    pen += w_clear × (clearance_radius − d)   # 가까울수록(d 작을수록) 큼, d=0 최대\n"
         "if w_tier:\n"
         "    pen += w_tier.get(cell.z, 0)               # 단 분리"),
        ("h2", "4.3 move_cost — 이동 1회 총 비용"),
        ("code",
         "c = cell_mm\n"
         "if prev_off is not None and move_off != prev_off:\n"
         "    c += w_turn                                # 방향 바뀜 = 회전\n"
         "c += cell_penalty(to_cell)\n"
         "return c"),
        ("h2", "4.4 astar_weighted 연동 — 방향을 가진 상태"),
        ("p", "turn penalty 는 '직전 진행 방향과 다른 방향으로 꺾을 때' 부과되므로, 같은 셀이라도 "
               "어느 방향으로 들어왔는지에 따라 이후 비용이 달라진다. 따라서 탐색 상태를 (셀, 진입방향)"
               "으로 확장한다. 진입방향은 NEIGHBORS_6 인덱스(0~5), 시작 상태는 −1(방향 없음). "
               "g / closed / came_from 모두 이 상태 단위로 관리한다(astar.py 의 astar_weighted)."),

        ("h1", "5. 자료구조 / 클래스"),
        ("table", (
            ["이름", "종류", "설명"],
            [
                ["RouteParams", "dataclass", "비용 가중치 묶음(아래 표). 음수면 ValueError."],
                ["CostModel", "class", "occ+params 로 이동비용/휴리스틱 계산. 클리어런스 맵 사전계산."],
                ["clearance", "np.int16[]", "CostModel 속성: 셀별 장애물까지 거리(없으면 None)."],
            ],
        )),

        ("h1", "6. 주요 함수"),
        ("table", (
            ["함수 / 메서드", "역할"],
            [
                ["clearance_map(occ, max_radius, connectivity=6)", "셀별 장애물 근접 거리맵(상한 max_radius)."],
                ["CostModel(occ, params)", "생성 시 클리어런스 맵 1회 사전계산."],
                ["CostModel.cell_penalty(cell)", "클리어런스 근접 + 단 분리 가산 페널티."],
                ["CostModel.move_cost(to, prev_off, move_off)", "이동 1회 총 비용(기본+회전+페널티)."],
                ["CostModel.heuristic(cell, goal)", "맨해튼×cell_mm. admissible & consistent."],
            ],
        )),

        ("h1", "7. 주요 변수 / 파라미터 (RouteParams)"),
        ("table", (
            ["필드", "기본값", "설명"],
            [
                ["cell_mm", "50.0", "셀 1칸 이동 기본 비용(mm)."],
                ["w_turn", "500.0", "회전 1회당 가산(mm, =셀 10칸). 클수록 회전 강하게 회피."],
                ["w_clear", "10.0", "클리어런스 페널티 계수(mm/셀). 0=비활성."],
                ["clearance_radius", "2", "페널티 적용 최대 근접 거리(셀). 이보다 멀면 0."],
                ["clearance_connectivity", "6", "거리 측정 이웃(6 맨해튼 / 26 체비셰프)."],
                ["w_tier", "{}", "단 분리 가중치 {z셀인덱스: 가산 mm}."],
            ],
        )),

        ("h1", "8. 실행 명령어"),
        ("code",
         "# 비용함수 적용 A* (회전·클리어런스 페널티)\n"
         ".\\.venv\\Scripts\\python.exe -m routing3d_py.astar ^\n"
         "    --region 195000 8000 14000 205000 12000 16000 --cell-mm 50 ^\n"
         "    --start 195300 8300 14775 --goal 204700 11700 14775 ^\n"
         "    --w-turn 500 --w-clear 10 --clearance 2 ^\n"
         "    --screenshot python_experiments/out/route_cost.png\n\n"
         "# 단위 테스트\n"
         ".\\.venv\\Scripts\\python.exe -m pytest python_experiments/tests/test_cost.py -v\n\n"
         "# 코드에서 사용\n"
         "from routing3d_py import astar_weighted, RouteParams\n"
         "res = astar_weighted(occ, start, goal,\n"
         "                     RouteParams(w_turn=500, w_clear=10, clearance_radius=2))"),
    ]
    render("Phase 1 · Step 1.3 — 비용함수(Cost Function)",
           "Routing3D · turn penalty / 클리어런스(가산 페널티) / 단 분리 · 모듈 cost.py",
           blocks, "step1_3_cost.docx")


# =============================================================================
# Step 1.4 — 다중 배관 순차 라우팅
# =============================================================================
def doc_multi():
    blocks = [
        ("h1", "1. 개요"),
        ("p", "여러 배관(start→end 작업)을 '한 개씩 차례로' 라우팅한다. 핵심은 이미 깔린 배관을 "
               "다음 배관의 장애물로 추가하여 배관끼리 같은 셀을 점유하지 않게(충돌 없이) 만드는 것이다. "
               "이것이 다중 배관 라우팅의 베이스라인(greedy sequential) 전략이다."),

        ("h1", "2. 순차 라우팅 흐름도"),
        ("code",
         "  1) 우선순위 규칙으로 작업 순서 결정 (기본: 긴 것 먼저)\n"
         "        │\n"
         "        ▼\n"
         "  2) 장애물 점유맵의 '작업용 사본' 생성 (원본 보존: occ.copy())\n"
         "        │\n"
         "        ▼  작업을 순서대로:\n"
         "     ┌───────────────────────────────────────────────────────────┐\n"
         "     │ a) start/end 가 점유면 가까운 빈 셀로 스냅(snap_to_free)   │\n"
         "     │ b) 비용함수 A*(astar_weighted)로 경로 탐색                 │\n"
         "     │ c) 성공 → 경로 셀(+pipe_radius 팽창)을 점유로 추가          │\n"
         "     │          → 이후 배관이 이 경로를 피한다                    │\n"
         "     │ d) 실패 → 기록만(다음 배관에 영향 없음)                    │\n"
         "     └───────────────────────────────────────────────────────────┘\n"
         "        │\n"
         "        ▼\n"
         "  4) 성공률 / 총 길이 / 실패 수 측정"),

        ("h1", "3. 충돌 회피 원리"),
        ("p", "각 배관을 라우팅한 직후 그 경로 셀을 작업용 점유맵에 '점유'로 표시(_mark_pipe)한다. "
               "다음 배관의 A* 는 is_blocked 로 이 셀들을 장애물처럼 피하므로, 성공한 경로끼리는 "
               "셀을 공유하지 않는다. pipe_radius>0 이면 경로 주변을 추가로 막아 배관 굵기/이격을 흉내낸다."),

        ("h1", "4. 우선순위 규칙(priority)"),
        ("table", (
            ["값", "정렬 기준"],
            [
                ["longest", "시작-끝 맨해튼 거리가 긴 배관 먼저(기본). 어려운 것 먼저."],
                ["shortest", "짧은 것 먼저."],
                ["utility", "유틸리티 라벨로 그룹(이름 순) 후, 그룹 내 긴 것 먼저."],
                ["original", "입력 순서 유지."],
            ],
        )),
        ("p", "계획서의 '직경 큰 순'은 직경 데이터 확보 시 추가 예정. 현재는 거리 기준."),

        ("h1", "5. 이 단계에서 하지 않는 것"),
        ("b", [
            "rip-up & reroute, CBS(충돌 기반 탐색) 등 전역 최적화 → Phase 3.",
            "따라서 혼잡한 출발부(메인장비 면에 PoC 밀집)에서는 후순위 배관이 막혀 실패할 수 있다. "
            "그 실패율을 '측정'하는 것이 본 단계의 목적이며, 해소는 Phase 3 과제다.",
        ]),

        ("h1", "6. 자료구조 / 클래스"),
        ("table", (
            ["이름", "종류", "설명"],
            [
                ["PipeResult", "dataclass", "배관 1개 결과: task / result(AStarResult) / order_index."],
                ["MultiRouteResult", "dataclass", "전체 결과: pipes / occupancy / priority + 지표 프로퍼티."],
            ],
        )),
        ("h2", "MultiRouteResult 지표"),
        ("table", (
            ["프로퍼티 / 메서드", "설명"],
            [
                ["success_count / fail_count", "성공·실패 배관 수."],
                ["total_length_mm", "성공 배관 기하 길이 합(mm)."],
                ["success_rate", "성공 비율(0~1)."],
                ["by_utility()", "결과를 유틸리티 라벨별로 묶은 dict."],
                ["summary()", "한 줄 요약 문자열."],
            ],
        )),

        ("h1", "7. 주요 함수"),
        ("table", (
            ["함수", "역할"],
            [
                ["order_tasks(occ, tasks, priority)", "우선순위 규칙으로 작업 정렬(원본 불변)."],
                ["route_sequential(occ, tasks, params, *, priority, pipe_radius, snap_to_free, max_expansions)",
                 "순차 라우팅 본체. occ 사본에 배관을 누적하며 충돌 없이 배치."],
                ["_mark_pipe(occ, path, radius)", "경로 셀(+반경 이웃)을 점유로 표시."],
                ["_snap(occ, cell, radius)", "점유 셀이면 반경 내 가장 가까운 빈 셀 반환."],
            ],
        )),

        ("h1", "8. 주요 변수 / 파라미터"),
        ("table", (
            ["이름", "기본값", "설명"],
            [
                ["priority", '"longest"', "작업 순서 규칙."],
                ["pipe_radius", "0", "깔린 배관을 점유로 추가할 때 팽창 반경(셀). 0=경로 셀만."],
                ["snap_to_free", "2", "start/end 가 점유면 빈 셀 탐색 반경(셀)."],
                ["max_expansions", "None", "배관당 A* 확장 상한(폭주 방지)."],
                ["work", "occ.copy()", "내부: 작업용 점유맵(장애물 + 누적 배관)."],
                ["ordered", "list[RouteTask]", "내부: 우선순위 정렬된 작업."],
            ],
        )),

        ("h1", "9. 실측 결과 — 프로젝트 6 (CLEAN/WTNHJ03)"),
        ("p", "장애물 983 · 메인장비 1 · 종단객체 57 · 라우팅 작업(PoC 페어) 208 · 유틸리티 21종. "
               "점유맵 셀 (126,136,81), 점유 256,928. cell_mm=100, priority=longest 로 전체를 충돌 없이 순차 라우팅."),
        ("p", "결과: 203 / 208 성공 (98%), 실패 5, 총 길이 3,677,400 mm. "
               "실패 5건은 혼잡 출발부 경합(H3PO4 3, NFW 2) — Step 1.4 가 '하지 않는 것'에서 예고한 "
               "후순위 경합 실패로, Phase 3 rip-up/CBS 의 대상이다."),
        ("table", (
            ["유틸리티", "성공/전체", "유틸리티", "성공/전체"],
            [
                ["[UPW] UPW_S", "42/42", "[Water] LPS", "5/5"],
                ["[Waste Liquid] NFW", "35/37", "[Water] LPR", "5/5"],
                ["[UPW] HOT DI_S", "36/36", "[Waste Water] FWW", "4/4"],
                ["[Gas] PA", "18/18", "[Exhaust] ALKA", "2/2"],
                ["[Exhaust] ACID", "17/17", "[Waste Liquid] T2(IPA)", "2/2"],
                ["[Water] NW", "10/10", "[Chemical] H3PO4", "5/8"],
                ["[Gas] PN2", "9/9", "기타 단일 유틸 7종", "각 1/1"],
                ["[Waste Water] NEU", "6/6", "합계", "203/208 (98%)"],
            ],
        )),

        ("h1", "10. 실행 명령어"),
        ("code",
         "# 프로젝트 6 전체 배관 충돌 없이 순차 라우팅 + 유틸리티별 렌더\n"
         ".\\.venv\\Scripts\\python.exe -m routing3d_py.scene --project 6 --multi ^\n"
         "    --priority longest --cell-mm 100 --screenshot python_experiments/out/multi.png\n\n"
         "# 단위 테스트\n"
         ".\\.venv\\Scripts\\python.exe -m pytest python_experiments/tests/test_multi_route.py -v\n\n"
         "# 코드에서 사용\n"
         "from routing3d_py import load_scene, route_sequential, RouteParams\n"
         "scene = load_scene(project_id=6)\n"
         "occ = scene.build_occupancy(cell_mm=100).occupancy\n"
         "mr = route_sequential(occ, scene.tasks, RouteParams(cell_mm=100, w_turn=300),\n"
         "                      priority='longest', pipe_radius=0)\n"
         "print(mr.summary())                  # 성공/실패/총길이/성공률\n"
         "for util, pipes in mr.by_utility().items():\n"
         "    ...                              # 유틸리티별 결과"),
    ]
    render("Phase 1 · Step 1.4 — 다중 배관 순차 라우팅",
           "Routing3D · greedy sequential(충돌 없이) · 우선순위 + 깔린 경로 점유 · 모듈 multi_route.py",
           blocks, "step1_4_multi_route.docx")


if __name__ == "__main__":
    doc_occupancy()
    doc_astar()
    doc_cost()
    doc_multi()
    print("done")
