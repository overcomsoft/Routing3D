# -*- coding: utf-8 -*-
r"""AI 자동설계 비교 리포트 기능 — 개발계획서(.docx) 생성기.

[이 스크립트가 하는 일]
  '메인장비별 · (장비,유틸리티그룹)별로 AI 자동설계(최단 / Stub+그룹패턴)를 수행하고,
   기존설계와 자동설계의 3D 뷰 + 배관 길이·꺾임·그룹핑 Factor 를 리포트로 출력'하는
   신규 기능의 개발계획서를 한글 .docx 로 만든다. 서식 헬퍼는 _gen_spec_docs.py 재사용.

[실행]  (프로젝트 루트에서)
  .\.venv\Scripts\python.exe python_experiments/out/_gen_autodesign_report_plan.py
  # 산출물: docs/routing3d_autodesign_report_plan.docx
  # PDF:    powershell -ExecutionPolicy Bypass -File python_experiments/out/_docx_to_pdf.ps1 `
  #           -in docs/routing3d_autodesign_report_plan.docx -out docs/routing3d_autodesign_report_plan.pdf
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _gen_spec_docs import (  # noqa: E402
    set_base_style, add_heading, add_para, add_bullets, add_code, add_table,
)
from docx import Document  # noqa: E402

OUT_PATH = os.path.normpath(
    os.path.join(os.path.dirname(__file__), "..", "..", "docs",
                 "routing3d_autodesign_report_plan.docx"))


def emit(doc, blocks):
    for kind, payload in blocks:
        if kind == "h1":
            add_heading(doc, payload, 1)
        elif kind == "h2":
            add_heading(doc, payload, 2)
        elif kind == "h3":
            add_heading(doc, payload, 3)
        elif kind == "p":
            add_para(doc, payload)
        elif kind == "b":
            add_bullets(doc, payload)
        elif kind == "code":
            add_code(doc, payload)
        elif kind == "table":
            add_table(doc, payload[0], payload[1])
        elif kind == "pb":
            doc.add_page_break()
        else:
            raise ValueError("unknown block kind: " + kind)


def add_cover(doc):
    add_heading(doc, "AI 자동설계 비교 리포트 — 개발계획서", 0)
    p = doc.add_paragraph()
    p.add_run(
        "메인장비별 · (장비, 유틸리티 그룹)별 AI 자동설계(최단 / Stub+그룹패턴)를 수행하고,\n"
        "기존설계와 자동설계의 3D 뷰 + 배관 길이·꺾임·그룹핑 Factor 를 리포트로 출력하는 기능").italic = True
    doc.add_paragraph("Routing3D · 단위 mm · DB DDW_AI_DB · C# WPF(HelixToolkit) + C++ 엔진(C ABI)")


# =============================================================================
# 1. 개요
# =============================================================================

def chap_overview():
    return [
        ("h1", "1. 개요"),
        ("h2", "1.1 목적"),
        ("p", "플랜트 한 동(툴 그룹)에 대해, 메인장비를 기준으로 (장비, 유틸리티 그룹) 단위로 "
              "AI 자동설계를 두 가지 전략으로 수행하고, 그 결과를 사람이 설계한 기존배관과 "
              "3차원 형상·정량 지표(길이·꺾임·그룹핑 Factor)로 한눈에 비교하는 리포트를 자동 "
              "생성한다. 목적은 (1) 자동설계 품질을 설계자가 신뢰·검수할 수 있게 가시화하고, "
              "(2) 전략(최단 vs Stub+그룹패턴)별 트레이드오프를 정량 근거로 제시하는 것이다."),
        ("h2", "1.2 배경"),
        ("p", "엔진은 이미 충돌 없는 직교 경로를 산출하고(다중배관 M1/M2), 기존설계 패턴 학습"
              "(스텁·랙·번들 L1~L4)을 자동 라우팅에 반영한다. 헤드리스 진단기 DbRouteDiag 는 "
              "성공률·총길이·꺾임·랙 집중도·번들 밀집도 등을 이미 측정하고, _gen_validation_report.py 는 "
              "기존설계 대비 길이비·꺾임·우회비를 docx/CSV 로 비교한다. 본 기능은 이 자산들을 "
              "'메인장비별 케이스 단위 + 3D 시각 비교 + 그룹핑 Factor'로 통합·확장하는 것이다."),
        ("h2", "1.3 범위(Scope)"),
        ("b", [
            "대상: 한 툴 그룹(TB_SPACE_GROUP_INFO) 내 메인장비(TB_EQUIPMENTS MAIN_SUB_TYPE='MainTool').",
            "케이스 단위: (메인장비) → (장비, 유틸리티 그룹) — 그룹 내 작업을 한 케이스로 묶어 라우팅·집계.",
            "자동설계 전략 2종: ① 최단(순수 A*) ② Stub+그룹패턴(스텁 고정 + 번들/랙 + 그룹코너).",
            "비교 대상: 기존설계(TB_ROUTE_PATH 폴리라인) vs 자동설계 2종.",
            "출력: 3D 뷰(기존/최단/Stub+그룹) + 지표 비교표(길이·꺾임·그룹핑 Factor) + 집계.",
            "비범위(1차): 실시간 편집·대화형 재설계(기존 뷰어 기능 유지), 운영 DB 쓰기.",
        ]),
        ("pb", None),
    ]


# =============================================================================
# 2. 기능 요구사항
# =============================================================================

def chap_requirements():
    return [
        ("h1", "2. 기능 요구사항"),
        ("h2", "2.1 입력"),
        ("table", (
            ["입력", "내용 · 출처"],
            [
                ["툴 그룹(프로젝트)", "TB_SPACE_GROUP_INFO 1건(또는 --all 전체 순회)"],
                ["메인장비 목록", "TB_EQUIPMENTS MAIN_SUB_TYPE='MainTool' (그룹 AABB 스코프)"],
                ["작업·기존설계", "TB_ROUTE_PATH(SOURCE_POS→TARGET_POS·폴리라인·BEND_COUNT·TOTAL_LENGTH)"],
                ["장애물/종단/PoC", "TB_BIM_OBSTACLE·TB_DUCT/LATERAL·TB_POCINSTANCES"],
                ["학습 자산", "TB_ROUTE_SEGMENT_TEMPLATE(스텁)·TB_ROUTE_DESIGN_GROUP(번들)"],
                ["실행 옵션", "cell_mm, 전략 선택, 출력 형식(docx/pdf/xlsx), 3D 스냅샷 on/off"],
            ],
        )),
        ("h2", "2.2 처리 — 케이스 매트릭스"),
        ("p", "메인장비마다, 그 장비에 연결된 작업을 유틸리티 그룹으로 묶어 케이스를 만든다. "
              "각 케이스에 대해 자동설계 2전략을 수행하고 기존설계와 함께 3-way 로 비교한다."),
        ("code",
         "for 메인장비 in 그룹.메인장비들:\n"
         "  for (장비, 유틸그룹) in 메인장비.케이스들:        # 케이스 = 같은 유틸그룹 작업 묶음\n"
         "     기존설계  = 기존배관 폴리라인(TB_ROUTE_PATH)    # ground truth\n"
         "     자동_최단 = route_multi(순수 A*, 패턴 OFF)       # 전략 A\n"
         "     자동_패턴 = route_multi(스텁ON+번들/랙+그룹코너)  # 전략 B\n"
         "     지표(기존/최단/패턴) = {길이, 꺾임, 그룹핑Factor}\n"
         "     3D스냅샷(기존/최단/패턴) = 오프스크린 렌더 PNG\n"
         "  케이스들 → 메인장비 섹션 → 그룹 집계 → 전체 집계"),
        ("h2", "2.3 자동설계 2전략 정의"),
        ("table", (
            ["전략", "구성(엔진 옵션) · 의도"],
            [
                ["A. 최단", "순수 weighted A*(동적 가중) · 패턴/번들/스텁 OFF. 비용최적·최단·직선적(기준선)."],
                ["B. Stub+그룹패턴", "스텁 라우팅 ON(고정 출발/종단 스텁) + 번들/랙(rack_levels·트렁크 z) + 그룹코너 경유. 사람 설계 추종·번들링 강화."],
            ],
        )),
        ("p", "전략은 기존 SceneViewModel 토글/환경변수(R3D_STUB·R3D_BUNDLE·R3D_RACK·R3D_GROUPCORNER·"
              "R3D_PATTERNS)와 1:1 매핑되며, 본 기능은 이를 케이스 배치로 자동 구성한다."),
        ("h2", "2.4 출력 — 리포트"),
        ("b", [
            "3D 뷰: 케이스별 기존/최단/Stub+그룹 3개 스냅샷(동일 카메라·축척, 유틸 색 일관).",
            "지표 비교표: 길이(총/평균)·꺾임(총/평균)·그룹핑 Factor(+세부) — 기존 대비 비율 포함.",
            "집계: (장비,유틸그룹)→메인장비→그룹 전체 롤업. 전략별 승패(짧음/꺾임적음/번들우수) 요약.",
            "형식: docx(3D 이미지 임베드) + PDF + xlsx(셀 단위 수치, 필터/피벗).",
        ]),
        ("pb", None),
    ]


# =============================================================================
# 3. 현재 자산 분석(재사용)
# =============================================================================

def chap_assets():
    return [
        ("h1", "3. 현재 자산 분석 (재사용 가능 모듈)"),
        ("p", "본 기능은 신규 알고리즘이 거의 없고, 기존 엔진·로더·진단·리포트 자산의 '오케스트레이션 + "
              "시각화'다. 아래 자산을 재사용해 개발량을 최소화한다."),
        ("table", (
            ["자산", "재사용 포인트"],
            [
                ["Engine(C ABI) / Interop", "route_multi·route_task·get_result·copy_path — 전략별 라우팅 실행"],
                ["ObstacleDbLoader", "그룹 스코프 로드(장애물·장비·작업·기존배관·PoC) + 격자 AABB 클램프"],
                ["SceneViewModel(라우팅 모드/토글)", "Shortest/GroupPattern·스텁·번들·랙·그룹코너 — 전략 구성 로직"],
                ["DbRouteDiag", "성공/총길이/turns/rackZ%/번들밀집%/lane/스텁매칭 측정 — 지표 산출 코어"],
                ["_gen_validation_report.py", "기존 vs 자동 길이비·꺾임·우회비 비교 + docx/CSV 골격"],
                ["ExistingPipe / StubExtractor", "기존설계 폴리라인·길이·꺾임 산출, 스텁 구간"],
                ["BundleStore / bundle_detect", "번들 멤버십·trunk z·pitch — 그룹핑 Factor 입력"],
                ["HelixToolkit 뷰어(BuildModel)", "3D 기하 구성 — 오프스크린 렌더로 스냅샷 캡처"],
                ["_gen_spec_docs.py", "docx 서식 헬퍼(제목/표/이미지) 재사용"],
            ],
        )),
        ("p", "신규 개발의 핵심은 (1) 케이스 매트릭스 배치 실행기, (2) 그룹핑 Factor 정식 정의·계산, "
              "(3) 오프스크린 3D 스냅샷 렌더러, (4) 3-way 비교 리포트 빌더 네 가지다."),
        ("pb", None),
    ]


# =============================================================================
# 4. 지표 정의
# =============================================================================

def chap_metrics():
    return [
        ("h1", "4. 지표 정의"),
        ("h2", "4.1 길이(Length)"),
        ("b", [
            "총 길이 = Σ 경로 기하 길이(mm). 자동=Σ(셀수−1)×cell_mm, 기존=폴리라인 누적 길이(또는 TOTAL_LENGTH).",
            "평균 길이 = 총 길이 / 성공 작업 수. 직선거리 대비 우회비 = 경로길이 / 종단점 직선거리.",
        ]),
        ("h2", "4.2 꺾임(Bends)"),
        ("b", [
            "꺾임 수 = 진행 축이 바뀌는 지점 수(90° 엘보). 자동=count_turns(path), 기존=BEND_COUNT 또는 dir-run 전환수.",
            "보조: 군더더기 꺾임(코너 대비 초과)·짧은 지그(staircase) — 경로 품질(톱니) 정량화.",
        ]),
        ("h2", "4.3 그룹핑 Factor (핵심 신규 지표)"),
        ("p", "여러 배관이 '함께 다발로 깔린 정도'를 0~1 로 정규화한 복합 지표. 사람 설계가 공용 랙을 "
              "따라 묶는 특성을 자동설계가 얼마나 재현했는지 측정한다. 하위 4개 성분의 가중 평균으로 정의하며, "
              "각 성분도 리포트에 개별 표기한다(해석 가능성)."),
        ("table", (
            ["성분", "정의 · 산출(기존 자산)", "방향"],
            [
                ["랙 집중도 (rackZ%)", "성공 경로의 수평 이동 셀 중 학습 랙 z-셀에 놓인 비율 (DbRouteDiag.Measure)", "높을수록 ↑"],
                ["번들 밀집도", "경로 쌍에서 한 배관 셀이 다른 배관 셀 ±R 안에 드는 평균 비율 (DbRouteDiag 번들밀집)", "높을수록 ↑"],
                ["Pitch 일관성", "인접 평행 배관 간격의 변동계수 CV → 일관성 = 1−min(1,CV) (bundle_detect pitch)", "높을수록 ↑"],
                ["레인 정렬도", "공용 트렁크 고도에 정렬된 작업 비율 = distinct lane 역수 기반 (AssignBundleLanes)", "높을수록 ↑"],
            ],
        )),
        ("code",
         "그룹핑Factor = w1·랙집중도 + w2·번들밀집도 + w3·pitch일관성 + w4·레인정렬도\n"
         "  (기본 가중 w = 0.35 / 0.30 / 0.20 / 0.15, 합 1.0 — 튜닝 가능)\n"
         "  각 성분은 [0,1] 정규화. 단일 배관 케이스는 그룹핑 정의상 N/A 표기."),
        ("p", "주: 가중치·반경 R·랙 z-셀 정의는 학습 자산(L3a/L4)과 동일 규약을 따른다. 기존설계의 "
              "그룹핑 Factor 도 같은 수식으로 계산해 자동설계와 동일 잣대로 비교한다."),
        ("pb", None),
    ]


# =============================================================================
# 5. 시스템 설계
# =============================================================================

def chap_design():
    return [
        ("h1", "5. 시스템 설계"),
        ("h2", "5.1 컴포넌트 구성"),
        ("code",
         "[배치 실행기 AutoDesignBatch]\n"
         "  ├─ CaseBuilder        : 그룹→메인장비→(장비,유틸그룹) 케이스 매트릭스 구성\n"
         "  ├─ StrategyRunner     : 케이스×전략(최단/Stub+그룹) 라우팅(Engine 재사용)\n"
         "  ├─ MetricsCalculator  : 길이·꺾임·그룹핑Factor(+성분) 산출(DbRouteDiag 코어 재사용)\n"
         "  ├─ SnapshotRenderer   : 오프스크린 HelixToolkit → 기존/자동 3D PNG\n"
         "  └─ ReportBuilder      : docx(이미지+표)+pdf+xlsx 생성\n"
         "[데이터 모델] DesignCase · DesignResult · MetricSet · CaseReport"),
        ("h2", "5.2 데이터 모델"),
        ("table", (
            ["타입", "필드(요지)"],
            [
                ["DesignCase", "MainEquip, Equip, UtilityGroup, TaskIndices[], ExistingPipes[]"],
                ["DesignResult", "Strategy(최단/패턴), Paths[], Success, LengthMm, Turns, Visited[]"],
                ["MetricSet", "TotalLen, AvgLen, TotalTurns, AvgTurns, GroupingFactor, {rackZ, density, pitchCV, lane}"],
                ["CaseReport", "DesignCase + {기존,최단,패턴} MetricSet + 3 SnapshotPaths"],
            ],
        )),
        ("h2", "5.3 3D 스냅샷 렌더링"),
        ("b", [
            "기존 BuildModel 기하 구성을 재사용하되, 화면 없이 렌더(Viewport3DHelper.RenderBitmap 또는 "
            "off-screen Viewport3D → RenderTargetBitmap → PNG).",
            "동일 카메라/축척/유틸 색으로 기존·최단·패턴 3장을 같은 시점에 렌더(공정 비교).",
            "케이스 AABB 에 ZoomExtents 후 캡처. 장애물 반투명·경로 튜브(유틸 색)·기존 폴리라인 표시.",
            "헤드리스 환경(서버) 대응: WPF 렌더는 STA 스레드 필요 → 배치 실행기를 STA 로 구동.",
        ]),
        ("h2", "5.4 실행 방식"),
        ("b", [
            "CLI 확장: Routing3D.Viewer.exe --autodesign-report <proj|ALL> <cell> <out_dir> [전략...] "
            "(기존 --dbroute 패턴 확장, 헤드리스).",
            "GUI 버튼: '자동설계 리포트' — 현재 그룹/선택 메인장비로 즉시 생성(진행 다이얼로그 재사용).",
            "산출물: out_dir/<그룹>_autodesign_report.{docx,pdf,xlsx} + snapshots/*.png.",
        ]),
        ("pb", None),
    ]


# =============================================================================
# 6. 리포트 구성
# =============================================================================

def chap_report():
    return [
        ("h1", "6. 리포트 구성(레이아웃)"),
        ("table", (
            ["섹션", "내용"],
            [
                ["표지", "그룹명·공정·일시·셀 크기·전략·엔진 버전"],
                ["요약", "케이스 수·성공률(전략별)·전체 길이/꺾임/그룹핑Factor 평균·전략 승패 요약"],
                ["메인장비별 섹션", "메인장비마다 (장비,유틸그룹) 케이스 목록"],
                ["케이스 상세", "3D 뷰 3개(기존/최단/Stub+그룹) + 지표 비교표 + 한줄 해석"],
                ["그룹 집계", "유틸그룹별 길이비·꺾임·그룹핑Factor 교차표"],
                ["부록", "실패 케이스 사유(혼잡/막힘)·파라미터·정의식"],
            ],
        )),
        ("h2", "6.1 케이스 상세 — 지표 비교표(예시 양식)"),
        ("table", (
            ["지표", "기존설계", "자동(최단)", "자동(Stub+그룹)"],
            [
                ["총 길이(mm)", "—", "—", "—"],
                ["평균 꺾임", "—", "—", "—"],
                ["그룹핑 Factor", "—", "—", "—"],
                ["  · 랙 집중도/번들밀집/pitch/레인", "—", "—", "—"],
                ["성공/작업", "(기존 N/A)", "ok/n", "ok/n"],
            ],
        )),
        ("p", "해석 규칙: '최단'은 길이·직선성 우위, 'Stub+그룹'은 그룹핑 Factor·사람설계 추종 우위가 "
              "기대치. 리포트는 케이스별로 어느 전략이 어느 지표에서 앞서는지 자동 표기한다."),
        ("pb", None),
    ]


# =============================================================================
# 7. 구현 단계
# =============================================================================

def chap_plan():
    return [
        ("h1", "7. 구현 단계(Phase)"),
        ("table", (
            ["Phase", "내용", "산출물 · 검증"],
            [
                ["P1 케이스 매트릭스", "CaseBuilder — 그룹→메인장비→(장비,유틸그룹) 케이스 구성. 기존 로더 재사용.", "케이스 목록 콘솔 출력 = DB 작업 수 일치"],
                ["P2 전략 실행기", "StrategyRunner — 최단/Stub+그룹 라우팅. SceneViewModel 토글 로직 추출·공유.", "--dbroute 결과와 케이스 합산 일치(회귀 0)"],
                ["P3 지표 산출", "MetricsCalculator — 길이·꺾임 + 그룹핑Factor(4성분) 정식화. DbRouteDiag 코어 분리·재사용.", "기존 vs 자동 수치 = 검증 리포트와 교차확인"],
                ["P4 3D 스냅샷", "SnapshotRenderer — 오프스크린 렌더(STA). 동일 카메라 3장.", "PNG 생성·시각 일관(유틸 색/축척)"],
                ["P5 리포트 빌더", "ReportBuilder — docx(이미지+표)+pdf+xlsx. _gen_spec_docs/_validation_report 재사용.", "1 그룹 리포트 산출·검수"],
                ["P6 통합·CLI/GUI", "--autodesign-report CLI + GUI 버튼 + 진행 다이얼로그.", "WTNHJ02 전 메인장비 리포트 e2e"],
                ["P7 집계·튜닝", "그룹/전체 롤업, 그룹핑Factor 가중 튜닝, --all 다중 그룹.", "다중 그룹 배치·성능(<목표) 확인"],
            ],
        )),
        ("h2", "7.1 일정(개략)"),
        ("b", [
            "P1~P3(코어 배치·지표): 약 1.5주 — 기존 자산 재사용으로 단축.",
            "P4(오프스크린 렌더): 약 1주 — WPF STA·카메라 정합이 리스크.",
            "P5~P6(리포트·통합): 약 1.5주.",
            "P7(집계·튜닝·다중그룹): 약 1주. 총 5주 내외(1인 기준, 버퍼 포함).",
        ]),
        ("pb", None),
    ]


# =============================================================================
# 8. 리스크 · 검증 · 확장
# =============================================================================

def chap_risk():
    return [
        ("h1", "8. 리스크 · 검증 · 향후 확장"),
        ("h2", "8.1 리스크 · 대응"),
        ("table", (
            ["리스크", "대응"],
            [
                ["오프스크린 WPF 렌더(헤드리스/서버)", "배치 실행기 STA 스레드 구동 · 실패 시 GUI 모드 캡처 폴백 · 또는 경량 3D(투영 PNG)"],
                ["그룹핑 Factor 가중 임의성", "성분 개별 표기 + 가중 외부 설정(env/인자) · 기존설계 동일식 비교로 상대화"],
                ["케이스 정의 모호(장비↔작업 매핑)", "TB_ROUTE_PATH SOURCE_GUID→PoC→장비 조인 규약 고정 · 미매핑 작업 '미상' 케이스로 격리"],
                ["혼잡/막힘 실패 케이스", "실패는 리포트에 사유(expanded>0=경로없음)로 명시 · rip-up/CBS 는 별도 과제"],
                ["다중 그룹 대량 렌더 비용", "스냅샷 on/off 옵션 · 해상도 조절 · xlsx 전용 모드(이미지 생략)"],
            ],
        )),
        ("h2", "8.2 검증 기준(수용 조건)"),
        ("b", [
            "케이스 합산 지표 = DbRouteDiag/검증 리포트 수치와 일치(회귀 0).",
            "동일 카메라로 기존/최단/패턴 3D 뷰가 시각적으로 비교 가능(유틸 색·축척 일관).",
            "그룹핑 Factor 가 'Stub+그룹 ≥ 최단' 경향(사람설계 추종) 재현 — 실데이터로 확인.",
            "WTNHJ02 전 메인장비 e2e 리포트(docx/pdf/xlsx) 무오류 생성.",
        ]),
        ("h2", "8.3 향후 확장"),
        ("b", [
            "전략 추가: FollowExisting(기존배관 복제)·rip-up·CBS 를 케이스 전략으로 편입.",
            "지표 추가: 자재 수량(엘보/티/플랜지) 산출 → 물량·원가 비교.",
            "인터랙티브 리포트: 웹/3D 뷰어 임베드(스냅샷→회전 가능 모델).",
            "기준선 학습 피드백: 그룹핑 Factor 낮은 케이스를 패턴 학습 보강 후보로 환류.",
        ]),
        ("p", "본 계획서는 신규 알고리즘보다 기존 자산(엔진·로더·진단·리포트·뷰어)의 통합·시각화에 "
              "초점을 둔다. 핵심 신규 산출물은 케이스 배치 실행기·그룹핑 Factor 정식화·오프스크린 "
              "3D 스냅샷·3-way 비교 리포트 빌더다."),
    ]


def build():
    doc = Document()
    set_base_style(doc)
    add_cover(doc)
    doc.add_page_break()
    for chap in (
        chap_overview, chap_requirements, chap_assets, chap_metrics,
        chap_design, chap_report, chap_plan, chap_risk,
    ):
        emit(doc, chap())
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print("저장 완료:", OUT_PATH)


if __name__ == "__main__":
    for st in (sys.stdout, sys.stderr):
        try:
            st.reconfigure(encoding="utf-8")
        except (AttributeError, ValueError):
            pass
    build()
